"""
generate_analysis_report.py
===========================
通用财务分析报告生成器。

输入:  extracted_data.json (由 collect_financial_deep.py 生成)
输出:  analysis_report.docx (带叙述分析的完整报告)

特性:
- 适配任意公司、任意年份范围
- LLM 将原始 JSON 归一化为标准 schema（亿元人民币）
- 数据缺口：Tavily 搜索 → LLM 估算 → 明确标注来源和置信度
- 每家公司单独一次 LLM 调用生成叙述段落
- 横向对比和核心结论由 LLM 综合生成
- 所有 LLM 调用经由 llm_client.py 的 4 模型 fallback 链

用法:
    python generate_analysis_report.py \\
        --data D:/clauderesult/claude0319/extracted_data.json \\
        --output D:/clauderesult/claude0319/ \\
        --query "分产品余额、收入和利润"
"""

import os
import sys
import json
import argparse
import re
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from llm_client import generate_content

try:
    from collect_search import tavily_search, tavily_extract
    TAVILY_AVAILABLE = True
except Exception:
    TAVILY_AVAILABLE = False

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


# ══════════════════════════════════════════════════════════════
# 1. 数据加载与 LLM 归一化
# ══════════════════════════════════════════════════════════════

def load_extracted_data(json_path: str) -> dict:
    with open(json_path, encoding='utf-8') as f:
        return json.load(f)


def normalize_company_data(company_name: str, raw: dict) -> dict:
    """
    用 LLM 将 collect_financial_deep.py 的原始 JSON 归一化为标准 schema。

    标准 schema:
    {
      "annual": {
        "2020": {
          "revenue": float | null,       # 亿元 RMB
          "net_income": float | null,    # 亿元 RMB
          "net_margin": float | null,    # 百分比, e.g. 25.3
          "loan_balance": float | null,  # 亿元 RMB（总贷款/促成余额）
          "total_assets": float | null   # 亿元 RMB
        },
        ...
      },
      "by_product": {
        "product_name": {"2020": float, "2021": float, ...},
        ...
      },
      "quarterly": {
        "2025-Q4": {
          "revenue": float | null,
          "net_income": float | null,
          "loan_balance": float | null,
          "by_product_revenue": {"product": float, ...}
        },
        ...
      },
      "semi_annual": {
        "2025-H1": {"revenue": float | null, "net_income": float | null, ...},
        ...
      },
      "unit": "亿元人民币",
      "data_source": str,
      "years_covered": [str, ...]
    }
    """
    if not raw or not raw.get('data'):
        return {"annual": {}, "by_product": {}, "quarterly": {}, "semi_annual": {},
                "unit": "亿元人民币", "data_source": raw.get('source', 'unknown'), "years_covered": []}

    # 将原始 data 序列化为简洁文本（截断过长内容）
    raw_text = json.dumps(raw['data'], ensure_ascii=False, indent=1)
    if len(raw_text) > 40000:
        raw_text = raw_text[:40000] + "\n... [截断] ..."

    source = raw.get('source', 'unknown')

    prompt = f"""你是金融数据分析专家。以下是从{company_name}({source})提取的原始财务数据JSON。
请将其归一化为标准格式。

【重要规则】
1. 所有金额统一换算为"亿元人民币"：
   - 若原始单位是"千元"(RMB Thousand)：÷ 100,000
   - 若原始单位是"百万元"：÷ 100
   - 若原始单位是"元"：÷ 100,000,000
   - 若已是亿元：直接使用
2. 期间格式：年度用"2024"，季度用"2024-Q4"，半年度用"2024-H1"
3. 找不到的数据填 null，禁止编造
4. by_product 中的产品名尽量用中文简称（如"微粒贷"、"信贷驱动服务"、"平台服务"）
5. loan_balance = 总贷款余额（含表内表外促成，取最大口径）
6. net_margin = net_income / revenue * 100（若两者都有则计算，否则null）

【原始数据】
{raw_text}

请直接输出 JSON（无需解释），格式：
{{
  "annual": {{
    "年份": {{"revenue": 数字或null, "net_income": 数字或null, "net_margin": 数字或null, "loan_balance": 数字或null, "total_assets": 数字或null}}
  }},
  "by_product": {{
    "产品名": {{"年份": 数字或null}}
  }},
  "quarterly": {{
    "期间": {{"revenue": 数字或null, "net_income": 数字或null, "loan_balance": 数字或null, "by_product_revenue": {{"产品": 数字或null}}}}
  }},
  "semi_annual": {{
    "期间": {{"revenue": 数字或null, "net_income": 数字或null, "loan_balance": 数字或null}}
  }},
  "unit": "亿元人民币",
  "data_source": "{source}",
  "years_covered": ["年份列表"]
}}"""

    resp = generate_content(prompt, max_output_tokens=16384, return_json=True)
    if isinstance(resp, dict):
        return resp

    # 尝试从文本提取 JSON
    try:
        m = re.search(r'\{[\s\S]+\}', resp or '')
        if m:
            return json.loads(m.group())
    except Exception:
        pass

    print(f"  ⚠ Normalization failed for {company_name}, returning empty schema")
    return {"annual": {}, "by_product": {}, "quarterly": {}, "semi_annual": {},
            "unit": "亿元人民币", "data_source": source, "years_covered": []}


# ══════════════════════════════════════════════════════════════
# 2. 数据缺口检测与搜索补全
# ══════════════════════════════════════════════════════════════

def detect_gaps(company_name: str, normalized: dict, target_years: list) -> list:
    """返回缺失年份列表（annual 中 revenue 和 net_income 都是 null 的年份）。"""
    annual = normalized.get('annual', {})
    gaps = []
    for y in target_years:
        yr = str(y)
        row = annual.get(yr, {})
        if row.get('revenue') is None and row.get('net_income') is None:
            gaps.append(yr)
    return gaps


def fill_gaps_with_search(company_name: str, normalized: dict, gaps: list, query: str) -> dict:
    """
    对缺失年份：Tavily 搜索 → LLM 估算 → 写入 normalized，并标注 estimate_sources。
    """
    if not gaps:
        return normalized

    print(f"    [Gap Fill] {company_name} missing: {gaps}")

    if not TAVILY_AVAILABLE:
        print(f"    ⚠ Tavily unavailable, skipping gap fill for {company_name}")
        return normalized

    # 构建搜索查询
    queries = [
        f"{company_name} annual revenue net profit {' '.join(gaps)} financial results",
        f"{company_name} {'  '.join(gaps)}年 营收 净利润 年度业绩",
    ]

    search_results = []
    for q in queries:
        try:
            results = tavily_search(q, max_results=4)
            search_results.extend(results)
            if len(search_results) >= 6:
                break
        except Exception as e:
            print(f"    ⚠ Search error: {e}")

    if not search_results:
        print(f"    ⚠ No search results for gap fill")
        return normalized

    # 全文提取 top 3
    urls = [r['url'] for r in search_results[:3] if r.get('url')]
    full_texts = []
    try:
        extracted = tavily_extract(urls)
        full_texts = [v for v in extracted.values() if v] if isinstance(extracted, dict) else []
    except Exception:
        full_texts = [r.get('content', '') for r in search_results[:3]]

    combined_text = '\n\n'.join(full_texts)[:15000]
    sources_used = [r.get('url', '') for r in search_results[:3]]

    prompt = f"""你是金融分析师。{company_name}在{', '.join(gaps)}年的财务数据缺失，请根据以下搜索资料进行合理估算。

【搜索资料】
{combined_text}

【现有数据（供参考/比对）】
{json.dumps(normalized.get('annual', {}), ensure_ascii=False, indent=1)[:3000]}

【搜索来源】
{sources_used}

【任务】
对缺失年份{gaps}，提供以下指标的估算值（亿元RMB）：
- revenue（营业收入）
- net_income（净利润）
- loan_balance（总贷款/信贷余额）
- confidence（"高"/"中"/"低"）

规则：
1. 只能基于搜索资料中出现的数字，不能凭空创造
2. 若搜索资料无任何线索，填 null
3. 必须写明该数字来自哪个来源（URL或机构名）
4. 估算值本质上是不确定的，confidence=低意味着可能误差超过50%

输出 JSON：
{{
  "年份": {{
    "revenue": 数字或null,
    "net_income": 数字或null,
    "loan_balance": 数字或null,
    "confidence": "高/中/低",
    "source": "来源描述"
  }}
}}"""

    resp = generate_content(prompt, max_output_tokens=2000, return_json=True)
    filled = resp if isinstance(resp, dict) else {}

    if not filled:
        try:
            m = re.search(r'\{[\s\S]+\}', str(resp))
            if m:
                filled = json.loads(m.group())
        except Exception:
            pass

    # 写入 normalized，标记为估算
    if not normalized.get('estimate_sources'):
        normalized['estimate_sources'] = {}

    for yr, vals in filled.items():
        if yr not in normalized['annual']:
            normalized['annual'][yr] = {}
        row = normalized['annual'][yr]
        for field in ['revenue', 'net_income', 'loan_balance']:
            if vals.get(field) is not None and row.get(field) is None:
                row[field] = vals[field]
                row[f'{field}_is_estimate'] = True

        confidence = vals.get('confidence', '低')
        source = vals.get('source', '网络搜索估算')
        normalized['estimate_sources'][yr] = {
            'confidence': confidence,
            'source': source,
            'search_urls': sources_used,
        }
        print(f"    ✓ Gap filled {yr}: revenue={vals.get('revenue')}, profit={vals.get('net_income')} (conf={confidence})")

    return normalized


# ══════════════════════════════════════════════════════════════
# 3. LLM 叙述生成
# ══════════════════════════════════════════════════════════════

ANALYST_STYLE = """你是一位资深金融科技行业分析师，撰写风格：
- 叙述性分析为主，数据支撑判断，不罗列数据
- 找出关键拐点并解释原因（如"2022年利润骤降，受消费信贷不良率上升及宏观逆风影响"）
- 对估算数据明确写"（估算，来源：XXX，置信度：X）"
- 不使用"前景广阔"、"潜力巨大"等空话
- 每段200-350字，有子标题"""


def generate_company_narrative(company_name: str, normalized: dict,
                                query: str, source_type: str) -> str:
    """生成单家公司的叙述分析文字（约600-1000字）。"""

    annual_str = json.dumps(normalized.get('annual', {}), ensure_ascii=False, indent=1)
    by_product_str = json.dumps(normalized.get('by_product', {}), ensure_ascii=False, indent=1)
    quarterly_str = json.dumps(normalized.get('quarterly', {}), ensure_ascii=False, indent=1)
    semi_str = json.dumps(normalized.get('semi_annual', {}), ensure_ascii=False, indent=1)
    est_str = json.dumps(normalized.get('estimate_sources', {}), ensure_ascii=False, indent=1)

    prompt = f"""{ANALYST_STYLE}

请为【{company_name}】撰写财务分析报告章节。

数据来源类型：{source_type}
研究重点：{query}
所有金额单位：亿元人民币

【年度数据（annual）】
{annual_str[:4000]}

【分产品数据（by_product）】
{by_product_str[:3000]}

【季度数据（quarterly）】
{quarterly_str[:3000]}

【半年度数据（semi_annual）】
{semi_str[:1000]}

【估算来源说明（estimate_sources）】
{est_str}

请撰写包含以下内容的分析文字（用###子标题分隔，不含表格）：
### 公司背景与数据说明
（说明公司定位、数据来源类型、数据完整度和任何重要局限）

### 收入趋势分析
（分析营收增长或下降的驱动力，找出拐点，提及分产品结构变化）

### 盈利能力分析
（净利润趋势、利润率变化、主要驱动/拖累因素）

### 信贷规模与产品结构
（贷款余额变化、分产品余额，如有估算明确标注）

### 最新高频数据洞察
（若有季度或半年度数据，重点分析最新趋势和与历史的对比；若无，则略去此节）

要求：
- 估算数字必须注明"（估算，来源：XXX）"
- 数据不足的维度写"受数据可得性限制，该项无法量化分析"
- 每个子节100-200字，有实质分析内容，不是数字罗列"""

    return generate_content(prompt, max_output_tokens=3000)


def generate_comparison_narrative(all_normalized: dict, query: str) -> str:
    """生成横向对比分析（约400-600字）。"""

    summary = {}
    for name, norm in all_normalized.items():
        ann = norm.get('annual', {})
        years = sorted(ann.keys())
        if years:
            latest = ann.get(years[-1], {})
            earliest = ann.get(years[0], {})
            summary[name] = {
                'years': years,
                'latest_year': years[-1],
                'latest_revenue': latest.get('revenue'),
                'latest_profit': latest.get('net_income'),
                'latest_loan': latest.get('loan_balance'),
                'earliest_revenue': earliest.get('revenue'),
                'revenue_cagr': _calc_cagr(earliest.get('revenue'), latest.get('revenue'), len(years) - 1),
                'profit_cagr': _calc_cagr(earliest.get('net_income'), latest.get('net_income'), len(years) - 1),
                'data_source': norm.get('data_source', 'unknown'),
            }

    prompt = f"""{ANALYST_STYLE}

请撰写以下公司的横向对比分析（400-600字），重点围绕：{query}

【各公司数据摘要（亿元人民币）】
{json.dumps(summary, ensure_ascii=False, indent=2)}

撰写要求：
1. 对比各公司的规模、盈利能力、增速差异
2. 解释差异的深层原因（商业模式、监管环境、战略选择）
3. 估算数据明确标注
4. 用###子标题：规模对比 / 盈利能力对比 / 增速对比 / 综合判断"""

    return generate_content(prompt, max_output_tokens=2000)


def generate_conclusion(all_normalized: dict, query: str, today: str) -> str:
    """生成核心结论速览（约250字，bullet 风格）。"""

    summaries = []
    for name, norm in all_normalized.items():
        ann = norm.get('annual', {})
        q = norm.get('quarterly', {})
        sa = norm.get('semi_annual', {})
        latest_yr = sorted(ann.keys())[-1] if ann else 'N/A'
        latest = ann.get(latest_yr, {})
        high_freq = {}
        if q:
            latest_q = sorted(q.keys())[-1]
            high_freq = {latest_q: q[latest_q]}
        if sa:
            latest_sa = sorted(sa.keys())[-1]
            high_freq[latest_sa] = sa[latest_sa]
        summaries.append({
            'company': name,
            'latest_annual': {latest_yr: latest},
            'high_freq': high_freq,
            'source': norm.get('data_source'),
        })

    prompt = f"""请为以下金融科技公司研究报告撰写"核心结论速览"（约250字）。
研究重点：{query}
数据截止：{today}

【数据摘要】
{json.dumps(summaries, ensure_ascii=False, indent=2)[:4000]}

格式要求：
- 1-2句总体概括（各公司整体路径差异）
- 每家公司1条bullet，含最新关键数字和最重要趋势判断
- 若有高频（季度/半年）数据，在对应公司bullet末尾补充
- 语气：分析师简报风格，直接有力"""

    return generate_content(prompt, max_output_tokens=800)


def generate_regulatory_background(companies: list, today: str) -> str:
    """生成监管背景与关键事件（直接由LLM生成，无需搜索，依靠训练数据）。"""
    prompt = f"""请为{', '.join(companies)}所在的中国消费金融/金融科技行业，
撰写2020-2025年关键监管事件年表（约300字）。

格式：表格形式（时间 | 事件 | 主要影响），4-8个最重要事件。
只写你确定的事实，不确定的不写。数据截止{today}。"""
    return generate_content(prompt, max_output_tokens=1000)


# ══════════════════════════════════════════════════════════════
# 4. Word 文档构建
# ══════════════════════════════════════════════════════════════

def _calc_cagr(start, end, years):
    try:
        if start and end and years > 0 and start > 0:
            return round((end / start) ** (1 / years) - 1, 4)
    except Exception:
        pass
    return None


def _fmt(val, decimals=1, suffix='亿'):
    if val is None:
        return '—'
    try:
        return f'{float(val):.{decimals}f}{suffix}'
    except Exception:
        return str(val)


def _fmt_pct(val):
    if val is None:
        return '—'
    try:
        return f'{float(val):.1f}%'
    except Exception:
        return str(val)


class ReportBuilder:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        s = self.doc.styles['Normal']
        s.font.name = 'Arial'
        s.font.size = Pt(11)
        s.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _run(self, para, text, bold=False, italic=False, size=11, color=None):
        run = para.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if color:
            run.font.color.rgb = RGBColor(*color)
        return run

    def heading(self, text, level=1):
        p = self.doc.add_heading(text, level=level)
        sz = {1: 16, 2: 14, 3: 12}.get(level, 11)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(sz)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        return p

    def para(self, text='', bold=False, italic=False):
        p = self.doc.add_paragraph()
        if text:
            self._run(p, text, bold=bold, italic=italic)
        return p

    def bullet(self, text):
        p = self.doc.add_paragraph(style='List Bullet')
        self._run(p, text)
        return p

    def narrative(self, text: str):
        """将 LLM 生成的 Markdown 文本（含 ### 子标题）写入文档。"""
        if not text:
            return
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith('### '):
                self.heading(stripped[4:], level=3)
            elif stripped.startswith('## '):
                self.heading(stripped[3:], level=2)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                self.bullet(stripped[2:])
            else:
                self.para(stripped)

    def shade_cell(self, cell, hex_color):
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def table(self, headers, rows, col_widths=None):
        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = 'Table Grid'
        hrow = t.rows[0]
        for i, h in enumerate(headers):
            cell = hrow.cells[i]
            self.shade_cell(cell, '1B3A5C')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(h))
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.rows[ri + 1].cells[ci]
                run = cell.paragraphs[0].add_run(str(val))
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in t.rows:
                    row.cells[i].width = Cm(w)
        return t

    def save(self, path):
        self.doc.save(path)
        print(f"  Saved: {path} ({os.path.getsize(path):,} bytes)")


# ══════════════════════════════════════════════════════════════
# 5. 表格构建（从归一化数据）
# ══════════════════════════════════════════════════════════════

def build_annual_table(builder: ReportBuilder, normalized: dict, company_name: str):
    """生成年度核心指标汇总表。"""
    ann = normalized.get('annual', {})
    est = normalized.get('estimate_sources', {})
    years = sorted(ann.keys())
    if not years:
        builder.para('（无年度数据）', italic=True)
        return

    headers = ['年份', '营收（亿元）', '净利润（亿元）', '净利润率', '贷款余额（亿元）', '数据类型']
    rows = []
    for yr in years:
        row = ann.get(yr, {})
        r_val = _fmt(row.get('revenue'))
        p_val = _fmt(row.get('net_income'))
        m_val = _fmt_pct(row.get('net_margin'))
        l_val = _fmt(row.get('loan_balance'))
        if yr in est:
            conf = est[yr].get('confidence', '低')
            dtype = f'估算（置信度：{conf}）'
        elif row.get('revenue_is_estimate') or row.get('net_income_is_estimate'):
            dtype = '部分估算'
        else:
            dtype = '年报提取'
        rows.append([yr, r_val, p_val, m_val, l_val, dtype])

    builder.table(headers, rows, col_widths=[2, 3.5, 3.5, 3, 4, 4])


def build_product_table(builder: ReportBuilder, normalized: dict):
    """生成分产品余额/收入表格。"""
    by_product = normalized.get('by_product', {})
    if not by_product:
        return

    # 收集所有年份
    all_years = sorted({yr for vals in by_product.values() for yr in vals.keys()})
    if not all_years:
        return

    headers = ['产品'] + all_years
    rows = []
    for product, year_vals in by_product.items():
        row = [product] + [_fmt(year_vals.get(yr)) for yr in all_years]
        rows.append(row)

    col_widths = [5] + [2.5] * len(all_years)
    builder.table(headers, rows, col_widths=col_widths)


def build_quarterly_table(builder: ReportBuilder, normalized: dict):
    """生成季度数据表格。"""
    quarterly = normalized.get('quarterly', {})
    if not quarterly:
        return

    periods = sorted(quarterly.keys())

    # 主要指标表
    headers = ['季度', '营收（亿元）', '净利润（亿元）', '贷款余额（亿元）']
    rows = []
    for p in periods:
        q = quarterly[p]
        rows.append([p, _fmt(q.get('revenue')), _fmt(q.get('net_income')), _fmt(q.get('loan_balance'))])
    builder.table(headers, rows, col_widths=[3, 3.5, 3.5, 4])
    builder.para()

    # 若有分产品收入，单独建表
    prod_data = {p: quarterly[p].get('by_product_revenue', {}) for p in periods
                 if quarterly[p].get('by_product_revenue')}
    if prod_data:
        all_products = sorted({prod for vals in prod_data.values() for prod in vals.keys()})
        if all_products:
            builder.para('分产品季度收入（亿元）：', bold=True)
            headers2 = ['产品'] + list(periods)
            rows2 = []
            for prod in all_products:
                row2 = [prod] + [_fmt(prod_data.get(p, {}).get(prod)) for p in periods]
                rows2.append(row2)
            col_widths2 = [6] + [2.5] * len(periods)
            builder.table(headers2, rows2, col_widths=col_widths2)


def build_semi_annual_table(builder: ReportBuilder, normalized: dict):
    """生成半年度数据表格。"""
    sa = normalized.get('semi_annual', {})
    if not sa:
        return
    periods = sorted(sa.keys())
    headers = ['期间', '营收（亿元）', '净利润（亿元）', '贷款余额（亿元）']
    rows = [[p, _fmt(sa[p].get('revenue')), _fmt(sa[p].get('net_income')),
             _fmt(sa[p].get('loan_balance'))] for p in periods]
    builder.table(headers, rows, col_widths=[3, 3.5, 3.5, 4])


def build_comparison_table(builder: ReportBuilder, all_normalized: dict):
    """生成横向对比汇总表。"""
    companies = list(all_normalized.keys())
    all_years = sorted({yr for norm in all_normalized.values()
                        for yr in norm.get('annual', {}).keys()})

    # 营收对比
    builder.para('营业收入对比（亿元人民币）', bold=True)
    rev_headers = ['年份'] + companies
    rev_rows = []
    for yr in all_years:
        row = [yr]
        for co in companies:
            ann = all_normalized[co].get('annual', {}).get(yr, {})
            val = _fmt(ann.get('revenue'))
            est = all_normalized[co].get('estimate_sources', {})
            if yr in est or ann.get('revenue_is_estimate'):
                val += '*'
            row.append(val)
        rev_rows.append(row)
    rev_rows.append(['*=估算数据'] + [''] * len(companies))
    builder.table(rev_headers, rev_rows)
    builder.para()

    # 净利润对比
    builder.para('净利润对比（亿元人民币）', bold=True)
    pft_rows = []
    for yr in all_years:
        row = [yr]
        for co in companies:
            ann = all_normalized[co].get('annual', {}).get(yr, {})
            val = _fmt(ann.get('net_income'))
            est = all_normalized[co].get('estimate_sources', {})
            if yr in est or ann.get('net_income_is_estimate'):
                val += '*'
            row.append(val)
        pft_rows.append(row)
    pft_rows.append(['*=估算数据'] + [''] * len(companies))
    builder.table(rev_headers, pft_rows)
    builder.para()

    # CAGR 对比
    builder.para('五年复合增长率（CAGR）对比', bold=True)
    cagr_headers = ['指标'] + companies
    cagr_rows = []
    for metric, label in [('revenue', '营收CAGR'), ('net_income', '净利润CAGR'), ('loan_balance', '贷款余额CAGR')]:
        row = [label]
        for co in companies:
            ann = all_normalized[co].get('annual', {})
            yrs = sorted(ann.keys())
            if len(yrs) >= 2:
                v = _calc_cagr(ann[yrs[0]].get(metric), ann[yrs[-1]].get(metric), len(yrs) - 1)
                row.append(_fmt_pct(v * 100 if v else None))
            else:
                row.append('—')
        cagr_rows.append(row)
    builder.table(cagr_headers, cagr_rows)


# ══════════════════════════════════════════════════════════════
# 6. 主流程
# ══════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description='通用财务分析报告生成器')
    parser.add_argument('--data', required=True, help='extracted_data.json 路径')
    parser.add_argument('--output', default='.', help='输出目录')
    parser.add_argument('--query', default='分产品收入、余额和利润', help='研究重点描述')
    parser.add_argument('--years', type=int, default=5, help='目标年数（用于缺口检测）')
    parser.add_argument('--name', default='financial_analysis', help='输出文件名前缀')
    args = parser.parse_args()

    today = datetime.now().strftime('%Y年%m月%d日')
    today_tag = datetime.now().strftime('%Y%m%d')

    print(f"\n{'='*60}")
    print(f"Financial Analysis Report Generator")
    print(f"  Data: {args.data}")
    print(f"  Query: {args.query}")
    print(f"  Date: {today}")
    print(f"{'='*60}\n")

    # 加载原始数据
    raw_data = load_extracted_data(args.data)
    companies = list(raw_data.keys())
    print(f"Companies: {companies}\n")

    # 确定目标年份范围
    current_year = datetime.now().year
    target_years = [str(current_year - i) for i in range(args.years, -1, -1)]

    # ── Step 1: 归一化 ──────────────────────────────────────
    print("Step 1: Normalizing data via LLM...")
    all_normalized = {}
    for company in companies:
        print(f"  [{company}] normalizing...")
        norm = normalize_company_data(company, raw_data.get(company, {}))
        all_normalized[company] = norm
        print(f"    Annual years: {sorted(norm.get('annual', {}).keys())}")
        print(f"    Quarterly periods: {sorted(norm.get('quarterly', {}).keys())}")

    # ── Step 2: 缺口检测 + 搜索补全 ─────────────────────────
    print("\nStep 2: Gap filling...")
    for company in companies:
        gaps = detect_gaps(company, all_normalized[company], target_years)
        if gaps:
            all_normalized[company] = fill_gaps_with_search(
                company, all_normalized[company], gaps, args.query)
        else:
            print(f"  [{company}] No gaps detected")

    # ── Step 3: LLM 叙述生成 ────────────────────────────────
    print("\nStep 3: Generating narratives via LLM...")
    narratives = {}
    for company in companies:
        print(f"  [{company}] writing narrative...")
        src = raw_data.get(company, {}).get('source', 'unknown')
        narratives[company] = generate_company_narrative(
            company, all_normalized[company], args.query, src)

    print("  Generating comparison...")
    comparison_text = generate_comparison_narrative(all_normalized, args.query)

    print("  Generating conclusion...")
    conclusion_text = generate_conclusion(all_normalized, args.query, today)

    print("  Generating regulatory background...")
    regulatory_text = generate_regulatory_background(companies, today)

    # ── Step 4: 组装 Word 文档 ───────────────────────────────
    print("\nStep 4: Building Word document...")
    b = ReportBuilder()

    # 标题
    title = b.doc.add_heading(
        f"{'、'.join(companies)} 财务数据对比报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(18)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    sub_p = b.doc.add_paragraph(f"研究重点：{args.query} | 数据截止：{today}")
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub_p.runs:
        run.font.italic = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    b.doc.add_paragraph()

    # 一、核心结论
    b.heading('一、核心结论速览', 2)
    b.narrative(conclusion_text)
    b.doc.add_paragraph()

    # 二~N、各公司章节
    for idx, company in enumerate(companies, 2):
        cn_num = ['二', '三', '四', '五', '六', '七', '八', '九', '十'][idx - 2] if idx <= 11 else str(idx)
        norm = all_normalized[company]
        src = raw_data.get(company, {}).get('source', 'unknown')

        b.heading(f'{cn_num}、{company}', 2)

        # 叙述分析（LLM）
        b.narrative(narratives[company])
        b.doc.add_paragraph()

        # 年度核心指标表
        b.heading('年度核心财务指标', 3)
        b.para('（亿元人民币；*=估算数据，来源见第末节数据说明）', italic=True)
        build_annual_table(b, norm, company)
        b.doc.add_paragraph()

        # 分产品表
        if norm.get('by_product'):
            b.heading('分产品数据', 3)
            build_product_table(b, norm)
            b.doc.add_paragraph()

        # 季度数据
        if norm.get('quarterly'):
            b.heading('最新季度数据', 3)
            build_quarterly_table(b, norm)
            b.doc.add_paragraph()

        # 半年度数据
        if norm.get('semi_annual'):
            b.heading('半年度数据', 3)
            build_semi_annual_table(b, norm)
            b.doc.add_paragraph()

    # 横向对比
    comp_num_idx = len(companies) + 1
    comp_cn = ['三', '四', '五', '六', '七', '八', '九', '十'][min(comp_num_idx - 2, 7)]
    b.heading(f'{comp_cn}、三家公司横向对比', 2)
    build_comparison_table(b, all_normalized)
    b.doc.add_paragraph()
    b.heading('综合分析', 3)
    b.narrative(comparison_text)
    b.doc.add_paragraph()

    # 监管背景
    next_idx = comp_num_idx + 1
    reg_cn = ['四', '五', '六', '七', '八', '九', '十'][min(next_idx - 3, 6)]
    b.heading(f'{reg_cn}、监管背景与关键事件', 2)
    b.narrative(regulatory_text)
    b.doc.add_paragraph()

    # 数据来源与说明
    last_cn = ['五', '六', '七', '八', '九', '十'][min(next_idx - 2, 5)]
    b.heading(f'{last_cn}、数据来源与说明', 2)

    b.heading('主要数据来源', 3)
    for company in companies:
        src = raw_data.get(company, {}).get('source', 'unknown')
        src_label = {
            'sec_edgar': 'SEC EDGAR 20-F年报 + 6-K EX-99.1季度新闻稿（自动提取）',
            'pdf_search': 'Tavily搜索年报PDF → pdfplumber提取',
            'web_search': '公开网络信息搜索（非官方财报，数字为估算）',
        }.get(src, src)
        b.bullet(f'{company}：{src_label}')
    b.bullet('提取工具：collect_financial_deep.py（SEC EDGAR API + LLM两轮提取）')
    b.bullet('分析报告：generate_analysis_report.py（LLM归一化 + 叙述生成）')
    b.bullet(f'报告生成日期：{today}')

    # 估算说明
    has_estimates = any(
        norm.get('estimate_sources') for norm in all_normalized.values()
    )
    if has_estimates:
        b.heading('估算数据说明', 3)
        for company in companies:
            est = all_normalized[company].get('estimate_sources', {})
            if est:
                b.para(f'{company}：以下年份数据为LLM搜索估算', bold=True)
                for yr, info in est.items():
                    conf = info.get('confidence', '低')
                    source = info.get('source', '网络搜索')
                    urls = info.get('search_urls', [])
                    url_str = '；'.join(urls[:2]) if urls else ''
                    b.bullet(f'{yr}年：置信度{conf}，来源：{source}' + (f'（{url_str}）' if url_str else ''))

    b.para('本报告不构成投资建议。估算数据存在较大不确定性，以官方披露为准。', italic=True)

    # 保存
    os.makedirs(args.output, exist_ok=True)
    out_name = args.name if args.name != 'financial_analysis' else f"financial_analysis_{today_tag}"
    out_path = os.path.join(args.output, f'{out_name}.docx')
    b.save(out_path)

    print(f"\n{'='*60}")
    print(f"DONE: {out_path}")
    print(f"{'='*60}\n")
    return out_path


if __name__ == '__main__':
    main()
