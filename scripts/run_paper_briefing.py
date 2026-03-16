"""
Paper Briefing — Phase 1
Fetches arXiv RSS → keyword filter → Word report.

Usage:
    python run_paper_briefing.py [--output output/] [--no-translate]

Output:
    - paper_briefing_YYYYMMDD.docx  (Word report)
    - paper_briefing_YYYYMMDD.json  (raw data)
"""

import os
import sys
import re
import json
import argparse
from datetime import datetime

# Add parent scripts dir to path
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SCRIPT_DIR)

from collect_rss import fetch_arxiv_rss
from utils import get_config_path


# ============================================================
# Config
# ============================================================

def load_briefing_config():
    """Load paper briefing configuration."""
    config_path = get_config_path("paper_briefing.json")
    if os.path.exists(config_path):
        with open(config_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {
        "focus_areas": {},
        "arxiv_categories": {
            "cs.AR": "计算机体系结构",
            "cs.AI": "人工智能",
            "cs.LG": "机器学习",
        },
        "hardware_spec_keywords": []
    }


# ============================================================
# Keyword Filtering
# ============================================================

def keyword_filter(papers, focus_areas):
    """
    Filter papers by keyword matching against title + abstract.
    Returns papers with added 'matched_area' and 'matched_keywords' fields.
    """
    filtered = []

    for paper in papers:
        searchable = f"{paper['title']} {paper['abstract']}".lower()
        matched = []

        for area, keywords in focus_areas.items():
            for kw in keywords:
                if kw.lower() in searchable:
                    matched.append((area, kw))

        if matched:
            paper['matched_area'] = matched[0][0]
            paper['matched_keywords'] = list(set(kw for _, kw in matched))
            filtered.append(paper)

    return filtered


# ============================================================
# Hardware Specs Extraction (from abstract, optional)
# ============================================================

def extract_hardware_specs(abstract, hw_keywords):
    """Check if abstract mentions hardware specs."""
    if not hw_keywords:
        return []
    text_lower = abstract.lower()
    return [kw for kw in hw_keywords if kw.lower() in text_lower]


# ============================================================
# Clean abstract (remove HTML tags + arXiv prefix from RSS)
# ============================================================

def clean_abstract(raw):
    """Remove HTML tags, arXiv prefix, and clean whitespace."""
    import html as html_lib
    text = re.sub(r'<[^>]+>', '', raw)
    text = html_lib.unescape(text)
    text = re.sub(r'\s+', ' ', text).strip()
    # Remove arXiv announcement header: "arXiv:XXXX.XXXXvN Announce Type: xxx Abstract: "
    text = re.sub(r'^arXiv:\d+\.\d+v\d+\s+Announce Type:\s*\S+\s+Abstract:\s*', '', text)
    # Remove plain "Abstract: " prefix
    if text.lower().startswith('abstract:'):
        text = text[9:].strip()
    return text


# ============================================================
# Batch Translate Abstracts via Gemini
# ============================================================

def translate_abstracts(papers, batch_size=5):
    """
    Translate paper abstracts from English to Chinese using Gemini Flash.
    Processes in batches for efficiency.
    Returns papers with added 'abstract_zh' field.
    """
    try:
        from llm_client import generate_content
    except ImportError:
        print("  [Warning] llm_client not available, skipping translation")
        return papers

    total = len(papers)
    translated = 0

    for i in range(0, total, batch_size):
        batch = papers[i:i+batch_size]
        # Build batch prompt
        abstracts_text = ""
        for j, paper in enumerate(batch):
            abstract = clean_abstract(paper.get('abstract', ''))
            abstracts_text += f"[{j+1}]\n{abstract}\n\n"

        prompt = f"""请将以下{len(batch)}篇学术论文的英文摘要准确翻译为中文。

翻译要求：
- 保持学术准确性，专业术语翻译准确
- 保留原文的专有名词首次出现时的英文（如 KV Cache、Processing-in-Memory 等）
- 语言风格与原文保持一致，不需要通俗化
- 每篇翻译之间用 [N] 标记分隔（N 为编号）

待翻译的摘要：

{abstracts_text}

请按 [1] [2] ... 的格式输出翻译结果，不要添加额外说明。"""

        try:
            result = generate_content(prompt, use_fast_model=True, temperature=0.1)
            # Parse translations by [N] markers
            parts = re.split(r'\[(\d+)\]', result)
            translations = {}
            for k in range(1, len(parts), 2):
                idx = int(parts[k])
                text = parts[k+1].strip() if k+1 < len(parts) else ''
                translations[idx] = text

            for j, paper in enumerate(batch):
                paper['abstract_zh'] = translations.get(j+1, '')
                if paper['abstract_zh']:
                    translated += 1

            print(f"  [翻译] 批次 {i//batch_size+1}: {len(translations)}/{len(batch)} 篇完成")
        except Exception as e:
            print(f"  [翻译错误] 批次 {i//batch_size+1}: {e}")
            for paper in batch:
                paper['abstract_zh'] = ''

    print(f"  [翻译完成] {translated}/{total} 篇成功翻译")
    return papers


# ============================================================
# Word Report Generation
# ============================================================

# Fixed area display order
AREA_ORDER = ['计算架构', '大模型优化', '系统与互联']
AREA_TITLES = {
    '计算架构': '计算架构类',
    '大模型优化': '大模型优化类',
    '系统与互联': '系统/互联类',
}

def generate_word_report(papers_by_area, output_path, config):
    """Generate a Word document briefing with Chinese labels."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # --- Helper: set font ---
    def _font(run, name_en='Arial', name_zh='SimSun', size=10.5, bold=False, color=None):
        run.font.name = name_en
        run.element.rPr.rFonts.set(qn('w:eastAsia'), name_zh)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

    # --- Title ---
    today = datetime.now().strftime('%Y年%m月%d日')
    title = doc.add_heading(f'学术论文追踪简报 — {today}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        _font(run, size=18, bold=True)

    # Summary stats
    total = sum(len(ps) for ps in papers_by_area.values())
    stats = doc.add_paragraph()
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = stats.add_run(f'共 {total} 篇匹配论文 | {len(papers_by_area)} 个研究方向')
    _font(r, size=10, color=(128, 128, 128))

    # --- Body: by focus area (fixed order) ---
    for area_key in AREA_ORDER:
        papers = papers_by_area.get(area_key, [])
        if not papers:
            continue

        area_title = AREA_TITLES.get(area_key, area_key)
        h = doc.add_heading(f'{area_title}（{len(papers)} 篇）', level=1)
        for run in h.runs:
            _font(run, size=14, bold=True)

        for i, paper in enumerate(papers, 1):
            # --- Title (English) ---
            p = doc.add_paragraph()
            r = p.add_run(f'{i}. {paper["title"]}')
            _font(r, size=11, bold=True)

            # --- 作者 (English names, Chinese label) ---
            if paper.get('authors'):
                p = doc.add_paragraph()
                r = p.add_run('作者：')
                _font(r, name_zh='SimSun', size=10, bold=True)
                r2 = p.add_run(paper['authors'])
                _font(r2, size=10, color=(80, 80, 80))

            # --- 链接 ---
            p = doc.add_paragraph()
            r = p.add_run('链接：')
            _font(r, name_zh='SimSun', size=10, bold=True)
            r2 = p.add_run(paper['link'])
            _font(r2, size=9, color=(0, 102, 204))

            # --- 匹配关键词 ---
            if paper.get('matched_keywords'):
                p = doc.add_paragraph()
                r = p.add_run('匹配关键词：')
                _font(r, name_zh='SimSun', size=9, bold=True, color=(160, 80, 0))
                r2 = p.add_run(', '.join(paper['matched_keywords']))
                _font(r2, size=9, color=(160, 80, 0))

            # --- 摘要 (Chinese translation, fallback to English) ---
            abstract_zh = paper.get('abstract_zh', '')
            abstract_en = clean_abstract(paper.get('abstract', ''))
            if abstract_zh:
                p = doc.add_paragraph()
                r = p.add_run('摘要：')
                _font(r, name_zh='SimSun', size=10, bold=True)
                r2 = p.add_run(abstract_zh)
                _font(r2, name_zh='SimSun', size=10)
            elif abstract_en:
                p = doc.add_paragraph()
                r = p.add_run('摘要（原文）：')
                _font(r, name_zh='SimSun', size=10, bold=True)
                r2 = p.add_run(abstract_en)
                _font(r2, size=10)

            # --- 硬件规格 (optional) ---
            hw_specs = paper.get('hardware_specs', [])
            if hw_specs:
                p = doc.add_paragraph()
                r = p.add_run('涉及硬件规格：')
                _font(r, name_zh='SimSun', size=9, bold=True, color=(0, 128, 0))
                r2 = p.add_run(', '.join(hw_specs))
                _font(r2, size=9, color=(0, 128, 0))

            # Separator
            if i < len(papers):
                sep = doc.add_paragraph()
                r = sep.add_run('─' * 60)
                _font(r, size=8, color=(200, 200, 200))

    doc.save(output_path)
    print(f"  [报告] 已保存: {output_path}")


# ============================================================
# Main Pipeline
# ============================================================

def main():
    parser = argparse.ArgumentParser(description='Paper Briefing — arXiv keyword tracker')
    parser.add_argument('--output', default=None, help='Output directory')
    parser.add_argument('--no-translate', action='store_true', help='Skip Gemini translation')
    args = parser.parse_args()

    # Output dir
    today_str = datetime.now().strftime('%m%d')
    if args.output:
        out_dir = args.output
    else:
        out_dir = os.path.join(f'D:\\clauderesult\\claude{today_str}', 'paper_briefing')
    os.makedirs(out_dir, exist_ok=True)

    # Load config
    config = load_briefing_config()
    focus_areas = config.get('focus_areas', {})
    categories = config.get('arxiv_categories', {})
    hw_keywords = config.get('hardware_spec_keywords', [])

    print("=" * 60)
    print("Step 1: 获取 arXiv RSS 论文")
    print("=" * 60)
    all_papers = fetch_arxiv_rss(categories=categories)

    print(f"\n{'=' * 60}")
    print("Step 2: 关键词过滤")
    print("=" * 60)
    print(f"  研究方向: {list(focus_areas.keys())}")
    all_keywords = [kw for kws in focus_areas.values() for kw in kws]
    print(f"  关键词数: {len(all_keywords)}")

    filtered = keyword_filter(all_papers, focus_areas)
    print(f"  匹配论文: {len(filtered)} / {len(all_papers)}")

    # Extract hardware specs
    for paper in filtered:
        paper['hardware_specs'] = extract_hardware_specs(
            paper.get('abstract', ''), hw_keywords
        )

    # Group by matched area
    papers_by_area = {}
    for paper in filtered:
        area = paper['matched_area']
        papers_by_area.setdefault(area, []).append(paper)

    for area, ps in papers_by_area.items():
        print(f"  [{area}] {len(ps)} 篇")

    # Step 3: Translate abstracts
    if not args.no_translate and filtered:
        print(f"\n{'=' * 60}")
        print("Step 3: 翻译摘要 (Gemini Flash)")
        print("=" * 60)
        translate_abstracts(filtered, batch_size=5)

    # Step 4: Generate reports
    print(f"\n{'=' * 60}")
    print("Step 4: 生成报告")
    print("=" * 60)

    date_str = datetime.now().strftime('%Y%m%d')

    # Save JSON
    json_path = os.path.join(out_dir, f'paper_briefing_{date_str}.json')
    json_data = []
    for paper in filtered:
        json_data.append({
            'title': paper['title'],
            'authors': paper.get('authors', ''),
            'abstract_en': clean_abstract(paper.get('abstract', '')),
            'abstract_zh': paper.get('abstract_zh', ''),
            'link': paper['link'],
            'arxiv_id': paper.get('arxiv_id', ''),
            'category': paper.get('category', ''),
            'matched_area': paper.get('matched_area', ''),
            'matched_keywords': paper.get('matched_keywords', []),
            'hardware_specs': paper.get('hardware_specs', []),
        })
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(json_data, f, ensure_ascii=False, indent=2)
    print(f"  [JSON] 已保存: {json_path}")

    # Generate Word
    if filtered:
        docx_path = os.path.join(out_dir, f'paper_briefing_{date_str}.docx')
        generate_word_report(papers_by_area, docx_path, config)
    else:
        print("  今日无匹配论文，未生成报告。")

    print(f"\n{'=' * 60}")
    print(f"完成 — {len(filtered)} 篇论文，{len(papers_by_area)} 个方向")
    print("=" * 60)


if __name__ == '__main__':
    main()
