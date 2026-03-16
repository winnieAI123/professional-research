"""
Domestic Policy Research Report — Word Report Generator
========================================================
Generates a formatted Word document from policy analysis data.
Handles markdown **bold** → Word bold conversion.
Supports markdown tables → Word tables.
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


FONT_CN = "宋体"
COLOR_HEADER = RGBColor(0x1B, 0x3A, 0x5C)
COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)
COLOR_ACCENT = RGBColor(0x25, 0x63, 0xEB)


def set_run(run, size=Pt(10), bold=False, color=None):
    """Apply font styling to a Word run."""
    run.font.name = "Arial"
    run.font.size = size
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    if color:
        run.font.color.rgb = color


def add_rich_paragraph(doc, text, base_size=Pt(10), is_bullet=False):
    """Add paragraph with **bold** markdown converted to real Word bold."""
    p = doc.add_paragraph(style="List Bullet") if is_bullet else doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run(run, size=base_size, bold=True)
        elif part:
            run = p.add_run(part)
            set_run(run, size=base_size)
    return p


def _render_markdown_table(doc, table_lines: list):
    """Convert markdown table lines into a formatted Word table.
    
    Expected input format:
    ['| Col1 | Col2 | Col3 |',
     '|------|------|------|',
     '| val1 | val2 | val3 |']
    """
    # Parse header and rows
    def parse_row(line):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        return [c for c in cells if c != '']
    
    rows = []
    for line in table_lines:
        line = line.strip()
        if re.match(r'^\|[\s\-:|]+\|$', line):  # separator line
            continue
        cells = parse_row(line)
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    # Determine column count from header
    num_cols = len(rows[0])
    
    # Create Word table
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    
    # Auto-fit columns
    table.autofit = True
    
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            
            # Remove any bold markers from cell text
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
            run = p.add_run(clean_text)
            
            if i == 0:  # Header row — just bold
                set_run(run, size=Pt(9), bold=True)
            else:
                set_run(run, size=Pt(9))
    
    # Add spacing after table
    doc.add_paragraph()


def generate_policy_report(
    report_text: str,
    analyses: list,
    output_dir: str,
    report_period: str = "",
    domain: str = "",
    focus_theme: str = "",
) -> str:
    """
    Generate a Word report from LLM-generated report text and analyses.

    Args:
        report_text: LLM output with [SECTION_*] markers
        analyses: List of structured analysis dicts
        output_dir: Directory to save the report
        report_period: e.g. "2026年3月"
        domain: e.g. "AI与机器人"
        focus_theme: e.g. "AI大模型监管与智算中心建设政策追踪"

    Returns:
        Path to generated .docx file
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    # ── Cover ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("政策分析月报")
    set_run(run, size=Pt(24), bold=True, color=COLOR_HEADER)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"{report_period} | {domain}")
    set_run(run2, size=Pt(12), color=COLOR_MUTED)

    if focus_theme:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(f"本期焦点: {focus_theme}")
        set_run(run3, size=Pt(11), bold=True, color=COLOR_ACCENT)

    doc.add_paragraph("─" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Sections ──
    sections = {
        "SECTION_COVER": "核心提要",
        "SECTION_1_MACRO": "一、宏观政策要览",
        "SECTION_2_DOMAIN": "二、分领域政策动态",
        "SECTION_3_LOCAL": "三、地方动态与机会扫描",
        "SECTION_4_DATA": "四、数据解读与市场风向",
        "SECTION_5_DEEP": "五、深度专题解读",
        "SECTION_6_STRATEGY": "六、业务影响与策略建议",
    }

    for marker, title in sections.items():
        # Extract section content
        pattern = rf'\[{marker}\](.*?)\[/{marker}\]'
        match = re.search(pattern, report_text, re.DOTALL)
        content = match.group(1).strip() if match else ""

        if not content:
            pattern2 = rf'{title}[：:](.*?)(?=\[SECTION_|\Z)'
            match2 = re.search(pattern2, report_text, re.DOTALL)
            content = match2.group(1).strip() if match2 else ""

        # Section heading
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"■ {title}")
        set_run(run, size=Pt(14), bold=True, color=COLOR_HEADER)

        # Section content
        if content:
            _render_section_content(doc, content)
        else:
            p = doc.add_paragraph()
            run = p.add_run("（本节内容待补充）")
            set_run(run, size=Pt(10), color=COLOR_MUTED)

    # ── Data Sources ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run("■ DATA SOURCES")
    set_run(run, size=Pt(12), bold=True, color=COLOR_HEADER)


    for a in analyses:
        t = a.get("title", "未知")
        u = a.get("url", "")
        p = doc.add_paragraph()
        run = p.add_run(f"• {t}: {u}")
        set_run(run, size=Pt(8), color=COLOR_MUTED)

    # ── Save ──
    period_code = re.sub(
        r'(\d{4})年(\d{1,2})月',
        lambda m: m.group(1) + m.group(2).zfill(2),
        report_period,
    )
    from datetime import datetime
    today = datetime.now().strftime("%Y%m%d")
    filepath = os.path.join(output_dir, f"policy_report_{period_code}_{today}.docx")

    # Avoid PermissionError if file is open
    if os.path.exists(filepath):
        try:
            doc.save(filepath)
        except PermissionError:
            filepath = filepath.replace(".docx", "_new.docx")
            doc.save(filepath)
    else:
        doc.save(filepath)

    print(f"  ✅ 报告生成: {filepath}")
    return filepath


def _render_section_content(doc, content: str):
    """Render markdown-like content into Word paragraphs with proper formatting.
    
    Layout hierarchy:
      Sub-domain header ("AI:") → bold, no indent
      Field lines ("- 政策名称:") → dash prefix, indent 0.5cm
      Numbered items ("1. xxx") → ■ marker, indent 1.0cm
      Normal paragraphs → indent 0.5cm
    """
    from docx.shared import Cm

    INDENT_FIELD = Cm(0.5)
    INDENT_NUMBERED = Cm(1.0)
    INDENT_BODY = Cm(0.5)

    def _is_subdomain_header(line: str) -> bool:
        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line).strip()
        if clean.endswith(":") or clean.endswith("："):
            if len(clean) <= 30 and not clean.startswith(("- ", "• ", "* ")):
                return True
        return False

    def _add_subdomain_header(doc, text: str):
        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text).strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(clean)
        set_run(run, size=Pt(11), bold=True, color=COLOR_HEADER)
        return p

    # Split into lines and process, collecting table blocks
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty/marker lines
        if not line or line.startswith("[") or line.startswith("/"):
            i += 1
            continue

        # Detect table block: consecutive lines starting with |
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) >= 2:  # header + separator at minimum
                _render_markdown_table(doc, table_lines)
            continue

        # Sub-domain header
        if _is_subdomain_header(line):
            _add_subdomain_header(doc, line)

        # Heading: ## or ###
        elif line.startswith("##"):
            clean = line.lstrip("# ").strip()
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', clean)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            run = p.add_run(clean)
            set_run(run, size=Pt(11), bold=True, color=COLOR_HEADER)

        # Numbered list: 1. 2. 3. → use ■ marker, deeper indent
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = INDENT_NUMBERED
            p.paragraph_format.space_after = Pt(2)
            # Replace "1." with "■ 1." for visual clarity
            text = re.sub(r'^(\d+\.)', r'■ \1', line)
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    set_run(run, size=Pt(9.5), bold=True)
                elif part:
                    run = p.add_run(part)
                    set_run(run, size=Pt(9.5))

        # Bullet list: - or • or * → field-level dash items
        elif re.match(r'^\s*[-•]\s+|^\s*\*\s+', line):
            clean_line = re.sub(r'^\s*[-•]\s+', '', line)
            clean_line = re.sub(r'^\s*\*\s+', '', clean_line)
            clean_line = clean_line.strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = INDENT_FIELD
            p.paragraph_format.space_after = Pt(2)
            # Add dash prefix
            parts = re.split(r'(\*\*.*?\*\*)', f"—  {clean_line}")
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    set_run(run, size=Pt(10), bold=True)
                elif part:
                    run = p.add_run(part)
                    set_run(run, size=Pt(10))

        # Pure bold line = sub-heading
        elif line.startswith("**") and line.endswith("**"):
            inner = line.strip("*").strip()
            if len(inner) <= 25:
                _add_subdomain_header(doc, line)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = INDENT_BODY
                run = p.add_run(inner)
                set_run(run, size=Pt(11), bold=True)

        # Normal paragraph with potential inline **bold**
        else:
            p = add_rich_paragraph(doc, line, base_size=Pt(10))
            p.paragraph_format.left_indent = INDENT_BODY

        i += 1


