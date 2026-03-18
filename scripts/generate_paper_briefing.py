# -*- coding: utf-8 -*-
"""
Paper Briefing Word Generator - Type 6 Academic Briefing

Generates a professionally formatted Word document from paper briefing data.
Matches the established format: centered title, colored keywords, 
proper horizontal separators, "摘要：" labels.

Usage:
    python generate_paper_briefing.py --input papers.json --output briefing.docx
    
Input JSON format:
{
    "date": "2026-03-18",
    "total_papers": 37,
    "total_directions": 3,
    "total_arxiv": 1824,
    "categories": [
        {
            "name": "计算架构",
            "papers": [
                {
                    "title": "English Paper Title",
                    "authors": "Author1, Author2",
                    "link": "https://arxiv.org/abs/xxxx",
                    "keywords": ["Neuromorphic", "PIM"],
                    "summary": "中文摘要...",
                    "hardware_specs": ["energy efficiency"]
                }
            ]
        }
    ],
    "blog_updates": [
        {
            "source": "Google AI",
            "articles": [
                {"title_zh": "中文标题", "summary": "一句话摘要", "link": "https://..."}
            ]
        }
    ]
}
"""

import argparse
import json
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Error: python-docx not installed. Run: python -m pip install python-docx")
    sys.exit(1)


# ============================================================
# STYLE CONFIGURATION
# ============================================================
STYLE = {
    'font_cn': '宋体',
    'font_en': 'Arial',
    'title_size': 18,
    'subtitle_size': 12,
    'heading_size': 14,
    'paper_title_size': 12,
    'body_size': 11,
    'label_size': 11,
    'small_size': 10,
    'color_title': RGBColor(0x1F, 0x4E, 0x79),      # Deep blue for title
    'color_heading': RGBColor(0x1F, 0x4E, 0x79),     # Deep blue for category headings
    'color_keyword': RGBColor(0x2E, 0x75, 0xB6),     # Medium blue for keywords
    'color_hwspec': RGBColor(0x2E, 0x75, 0xB6),      # Medium blue for hardware specs
    'color_link': RGBColor(0x05, 0x63, 0xC1),        # Link blue
    'color_label': RGBColor(0x40, 0x40, 0x40),        # Dark gray for labels
    'color_body': RGBColor(0x33, 0x33, 0x33),         # Body text
    'separator_color': 'A0A0A0',                       # Gray for separator lines
}


def _set_run_font(run, size=None, bold=False, italic=False, color=None):
    """Apply font settings to a run."""
    size = size or STYLE['body_size']
    run.font.name = STYLE['font_en']
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    run.element.rPr.rFonts.set(qn('w:eastAsia'), STYLE['font_cn'])


def _add_horizontal_line(doc, color=None):
    """Add a proper Word horizontal line (border on paragraph bottom)."""
    color = color or STYLE['separator_color']
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # Create bottom border using XML
    pPr = p._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)


def _add_labeled_line(doc, label, value, label_color=None, value_color=None, 
                      value_bold=False, label_size=None, value_size=None):
    """Add a line like '作者：John Doe' with separate formatting for label and value."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    # Label run
    run_label = p.add_run(f'{label}')
    _set_run_font(run_label, size=label_size or STYLE['label_size'], 
                  bold=True, color=label_color or STYLE['color_label'])
    
    # Value run
    run_value = p.add_run(f'{value}')
    _set_run_font(run_value, size=value_size or STYLE['body_size'],
                  bold=value_bold, color=value_color or STYLE['color_body'])
    
    return p


def generate_paper_briefing_word(data, output_path):
    """
    Generate professionally formatted Word document for paper briefing.
    
    Args:
        data: dict with paper briefing data (see module docstring for format)
        output_path: path to save .docx file
    """
    doc = Document()
    
    # --- Set default style ---
    style = doc.styles['Normal']
    style.font.name = STYLE['font_en']
    style.font.size = Pt(STYLE['body_size'])
    style.element.rPr.rFonts.set(qn('w:eastAsia'), STYLE['font_cn'])
    
    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ========================================
    # TITLE
    # ========================================
    date_str = data.get('date', datetime.now().strftime('%Y-%m-%d'))
    # Convert YYYY-MM-DD to YYYY 年 MM 月 DD 日
    try:
        dt = datetime.strptime(date_str, '%Y-%m-%d')
        date_display = f'{dt.year} 年 {dt.month:02d} 月 {dt.day:02d} 日'
    except ValueError:
        date_display = date_str
    
    title_text = f'学术论文追踪简报 — {date_display}'
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(6)
    run = p_title.add_run(title_text)
    _set_run_font(run, size=STYLE['title_size'], bold=True, color=STYLE['color_title'])
    
    # ========================================
    # SUBTITLE (summary line)
    # ========================================
    total = data.get('total_papers', 0)
    directions = data.get('total_directions', 0)
    total_arxiv = data.get('total_arxiv', 0)
    
    subtitle_parts = [f'共 {total} 篇匹配论文', f'{directions} 个研究方向']
    subtitle_text = ' | '.join(subtitle_parts)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(12)
    run = p_sub.add_run(subtitle_text)
    _set_run_font(run, size=STYLE['subtitle_size'], color=STYLE['color_body'])
    
    # Data source description (if available)
    if total_arxiv:
        p_source = doc.add_paragraph()
        p_source.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_source.paragraph_format.space_after = Pt(6)
        source_text = f'数据来源：arXiv RSS 当日新论文（共 {total_arxiv} 篇），经关键词预筛选 + LLM 精选'
        run = p_source.add_run(source_text)
        _set_run_font(run, size=STYLE['small_size'], color=RGBColor(0x80, 0x80, 0x80))
    
    # ========================================
    # CATEGORIES & PAPERS
    # ========================================
    categories = data.get('categories', [])
    
    for cat in categories:
        cat_name = cat.get('name', '未分类')
        papers = cat.get('papers', [])
        
        if not papers:
            continue
        
        # Category heading
        p_cat = doc.add_paragraph()
        p_cat.paragraph_format.space_before = Pt(18)
        p_cat.paragraph_format.space_after = Pt(8)
        run = p_cat.add_run(f'{cat_name}类（{len(papers)} 篇）')
        _set_run_font(run, size=STYLE['heading_size'], bold=True, color=STYLE['color_heading'])
        
        for idx, paper in enumerate(papers):
            # Paper number + title (bold)
            p_title_line = doc.add_paragraph()
            p_title_line.paragraph_format.space_before = Pt(8)
            p_title_line.paragraph_format.space_after = Pt(4)
            run = p_title_line.add_run(f'{idx + 1}. {paper.get("title", "Untitled")}')
            _set_run_font(run, size=STYLE['paper_title_size'], bold=True, color=STYLE['color_heading'])
            
            # Authors
            authors = paper.get('authors', '')
            if authors:
                _add_labeled_line(doc, '作者：', authors)
            
            # Link
            link = paper.get('link', '')
            if link:
                _add_labeled_line(doc, '链接：', link, value_color=STYLE['color_link'])
            
            # Matched keywords (blue color)
            keywords = paper.get('keywords', [])
            if keywords:
                if isinstance(keywords, list):
                    keywords_text = ', '.join(keywords)
                else:
                    keywords_text = str(keywords)
                _add_labeled_line(doc, '匹配关键词：', keywords_text, 
                                  label_color=STYLE['color_keyword'],
                                  value_color=STYLE['color_keyword'])
            
            # Summary with "摘要：" label
            summary = paper.get('summary', '')
            if summary:
                p_summary = doc.add_paragraph()
                p_summary.paragraph_format.space_before = Pt(6)
                p_summary.paragraph_format.space_after = Pt(4)
                
                # "摘要：" label
                run_label = p_summary.add_run('摘要：')
                _set_run_font(run_label, bold=True, color=STYLE['color_label'])
                
                # Summary text
                run_text = p_summary.add_run(summary)
                _set_run_font(run_text, color=STYLE['color_body'])
            
            # Hardware specs (blue, if any)
            hw_specs = paper.get('hardware_specs', [])
            if hw_specs:
                if isinstance(hw_specs, list):
                    hw_text = ', '.join(hw_specs)
                else:
                    hw_text = str(hw_specs)
                if hw_text and hw_text.lower() not in ('', 'none', '[]', 'null'):
                    _add_labeled_line(doc, '涉及硬件规格：', hw_text,
                                      label_color=STYLE['color_hwspec'],
                                      value_color=STYLE['color_hwspec'])
            
            # Horizontal separator between papers
            _add_horizontal_line(doc)
    
    # ========================================
    # BLOG UPDATES
    # ========================================
    blog_updates = data.get('blog_updates', [])
    if blog_updates:
        p_blog_heading = doc.add_paragraph()
        p_blog_heading.paragraph_format.space_before = Pt(24)
        p_blog_heading.paragraph_format.space_after = Pt(8)
        run = p_blog_heading.add_run('AI Lab 博客更新')
        _set_run_font(run, size=STYLE['heading_size'], bold=True, color=STYLE['color_heading'])
        
        for source_group in blog_updates:
            source_name = source_group.get('source', '')
            articles = source_group.get('articles', [])
            
            if not articles:
                continue
            
            # Source name as sub-heading
            p_source = doc.add_paragraph()
            p_source.paragraph_format.space_before = Pt(8)
            run = p_source.add_run(source_name)
            _set_run_font(run, size=STYLE['body_size'], bold=True, color=STYLE['color_heading'])
            
            for article in articles:
                title_zh = article.get('title_zh', article.get('title', ''))
                summary = article.get('summary', '')
                link = article.get('link', '')
                
                # Format: title | summary | link
                p_article = doc.add_paragraph()
                p_article.paragraph_format.left_indent = Inches(0.3)
                p_article.paragraph_format.space_before = Pt(2)
                p_article.paragraph_format.space_after = Pt(2)
                
                # Bullet
                run_bullet = p_article.add_run('• ')
                _set_run_font(run_bullet)
                
                # Title
                run_title = p_article.add_run(title_zh)
                _set_run_font(run_title, bold=True, color=STYLE['color_body'])
                
                if summary:
                    run_sep = p_article.add_run(' | ')
                    _set_run_font(run_sep, color=RGBColor(0x99, 0x99, 0x99))
                    run_sum = p_article.add_run(summary)
                    _set_run_font(run_sum, color=STYLE['color_body'])
                
                if link:
                    run_sep = p_article.add_run(' | ')
                    _set_run_font(run_sep, color=RGBColor(0x99, 0x99, 0x99))
                    run_link = p_article.add_run(link)
                    _set_run_font(run_link, size=STYLE['small_size'], color=STYLE['color_link'])
    
    # ========================================
    # FOOTER
    # ========================================
    _add_horizontal_line(doc)
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    now = datetime.now().strftime('%Y年%m月%d日 %H:%M')
    run = p_footer.add_run(f'简报生成时间：{now}')
    _set_run_font(run, size=STYLE['small_size'], color=RGBColor(0x99, 0x99, 0x99))
    
    # Save
    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
    doc.save(output_path)
    print(f'  ✓ Word report saved: {output_path}')
    print(f'    Size: {os.path.getsize(output_path):,} bytes')
    
    return output_path


# ============================================================
# CLI INTERFACE
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description='Generate Paper Briefing Word Report (Type 6)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
    python generate_paper_briefing.py --input papers.json --output briefing.docx
    python generate_paper_briefing.py --input papers.json  # Output: same dir, .docx
'''
    )
    parser.add_argument('--input', '-i', required=True, help='Input JSON file with paper data')
    parser.add_argument('--output', '-o', help='Output Word file path (default: same dir, .docx)')
    
    args = parser.parse_args()
    
    # Load JSON
    with open(args.input, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Determine output path
    if args.output:
        output_path = args.output
    else:
        base = os.path.splitext(args.input)[0]
        output_path = base + '.docx'
    
    generate_paper_briefing_word(data, output_path)


if __name__ == '__main__':
    main()
