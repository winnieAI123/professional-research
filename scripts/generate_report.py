"""
Report generation: Markdown to Word (.docx) conversion.
Preserves headings, tables, lists, bold, italic, and links.
Chinese font: SimSun (宋体), English font: Arial.
"""

import os
import re

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("  [Error] python-docx not installed. Run: python -m pip install python-docx")

from utils import get_output_dir, sanitize_filename


# ============================================================
# Font Configuration
# ============================================================

FONT_CHINESE = "SimSun"      # 宋体
FONT_ENGLISH = "Arial"
FONT_SIZE_BODY = 10.5         # 五号
FONT_SIZE_H1 = 18
FONT_SIZE_H2 = 15
FONT_SIZE_H3 = 12
FONT_SIZE_H4 = 11
LINE_SPACING = 1.5


# ============================================================
# Core Conversion
# ============================================================

def _set_run_font(run, bold=False, italic=False, size=None, color=None):
    """Set font properties for a run element."""
    run.font.name = FONT_ENGLISH
    run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CHINESE)
    
    if size:
        run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_formatted_text(paragraph, text, bold=False, italic=False,
                        size=FONT_SIZE_BODY, color=None):
    """Add text with formatting to a paragraph."""
    # Handle inline formatting: **bold** and *italic*
    if not bold and not italic and ("**" in text or "*" in text):
        _add_inline_formatted_text(paragraph, text, size, color)
        return
    
    run = paragraph.add_run(text)
    _set_run_font(run, bold=bold, italic=italic, size=size, color=color)


def _add_inline_formatted_text(paragraph, text, size=FONT_SIZE_BODY, color=None):
    """Parse and render inline markdown formatting (bold, italic)."""
    # Pattern: **bold**, *italic*, ***bold+italic***
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before match
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            _set_run_font(run, size=size, color=color)
        
        if match.group(2):  # ***bold+italic***
            run = paragraph.add_run(match.group(2))
            _set_run_font(run, bold=True, italic=True, size=size, color=color)
        elif match.group(3):  # **bold**
            run = paragraph.add_run(match.group(3))
            _set_run_font(run, bold=True, size=size, color=color)
        elif match.group(4):  # *italic*
            run = paragraph.add_run(match.group(4))
            _set_run_font(run, italic=True, size=size, color=color)
        
        last_end = match.end()
    
    # Add remaining text
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        _set_run_font(run, size=size, color=color)


def _set_paragraph_spacing(paragraph):
    """Set standard paragraph spacing."""
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(3)
    fmt.space_after = Pt(3)
    fmt.line_spacing = LINE_SPACING


def markdown_to_docx(md_text: str, output_path: str) -> str:
    """
    Convert Markdown text to Word document.
    
    Supports:
    - Headings (# to ####)
    - Tables (| col1 | col2 |)
    - Unordered lists (- or *)
    - Ordered lists (1. 2. 3.)
    - Bold (**text**) and Italic (*text*)
    - Horizontal rules (---)
    - Blockquotes (> text)
    
    Args:
        md_text: Markdown content string
        output_path: Full path for the output .docx file
    
    Returns:
        Absolute path of the generated Word file
    """
    doc = Document()
    
    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_ENGLISH
    font.size = Pt(FONT_SIZE_BODY)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CHINESE)
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    lines = md_text.split("\n")
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            i += 1
            continue
        
        # --- Horizontal Rule ---
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            # Add a thin horizontal line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # --- Headings ---
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            
            sizes = {1: FONT_SIZE_H1, 2: FONT_SIZE_H2,
                     3: FONT_SIZE_H3, 4: FONT_SIZE_H4}
            
            p = doc.add_paragraph()
            _add_formatted_text(p, text, bold=True,
                              size=sizes.get(level, FONT_SIZE_BODY),
                              color=(26, 54, 93))
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # --- Table ---
        if "|" in stripped and stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and "|" in lines[i].strip():
                table_lines.append(lines[i].strip())
                i += 1
            
            _add_table(doc, table_lines)
            continue
        
        # --- Blockquote ---
        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            # Skip alert markers like [!NOTE] [!IMPORTANT]
            if text.startswith("[!"):
                i += 1
                continue
            p = doc.add_paragraph()
            _add_formatted_text(p, text, italic=True,
                              color=(100, 100, 100))
            _set_paragraph_spacing(p)
            p.paragraph_format.left_indent = Cm(1)
            i += 1
            continue
        
        # --- Unordered List ---
        if re.match(r'^[-*]\s+', stripped):
            text = re.sub(r'^[-*]\s+', '', stripped)
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)
            _set_paragraph_spacing(p)
            i += 1
            continue
        
        # --- Ordered List ---
        if re.match(r'^\d+\.\s+', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, text)
            _set_paragraph_spacing(p)
            i += 1
            continue
        
        # --- Regular Paragraph ---
        p = doc.add_paragraph()
        _add_formatted_text(p, stripped)
        _set_paragraph_spacing(p)
        i += 1
    
    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"  [Word] Report saved: {output_path}")
    return output_path


def _add_table(doc, table_lines: list):
    """Parse markdown table lines and add to document."""
    if len(table_lines) < 2:
        return
    
    # Parse header
    headers = [c.strip() for c in table_lines[0].split("|") if c.strip()]
    
    # Skip separator line (|---|---|)
    data_start = 1
    if len(table_lines) > 1 and re.match(r'^[\|\s\-:]+$', table_lines[1]):
        data_start = 2
    
    # Parse data rows
    rows = []
    for line in table_lines[data_start:]:
        cells = [c.strip() for c in line.split("|") if c.strip()]
        rows.append(cells)
    
    if not headers:
        return
    
    # Create table
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Fill header
    for j, header in enumerate(headers):
        if j < num_cols:
            cell = table.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_formatted_text(p, header, bold=True, size=9,
                              color=(255, 255, 255))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Header background color
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="1A365D" w:val="clear"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)
    
    # Fill data rows
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.rows[i + 1].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                _add_formatted_text(p, cell_text, size=9)
                
                # Alternating row color
                if i % 2 == 0:
                    from docx.oxml.ns import nsdecls
                    from docx.oxml import parse_xml
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="F0F4F8" w:val="clear"/>'
                    )
                    cell._tc.get_or_add_tcPr().append(shading)
    
    # Add spacing after table
    doc.add_paragraph()


# ============================================================
# Convenience: Save both MD and Word
# ============================================================

def save_report(
    md_content: str,
    topic: str,
    output_dir: str = None,
) -> dict:
    """
    Save report in both Markdown and Word formats.
    
    Args:
        md_content: The markdown report content
        topic: Research topic (used for filename)
        output_dir: Output directory. If None, uses default workspace.
    
    Returns:
        Dict with 'md_path' and 'docx_path'
    """
    if output_dir is None:
        output_dir = get_output_dir()
    
    os.makedirs(output_dir, exist_ok=True)
    safe_name = sanitize_filename(topic)
    
    # Save MD
    md_path = os.path.join(output_dir, f"{safe_name}_report.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    print(f"  [MD] Report saved: {md_path}")
    
    # Save Word
    docx_path = os.path.join(output_dir, f"{safe_name}_report.docx")
    markdown_to_docx(md_content, docx_path)
    
    return {
        "md_path": md_path,
        "docx_path": docx_path,
    }
