"""
Multi-Source Leaderboard Word Report
=====================================
Generates professional Word report with embedded charts.
Adapted for professional-research skill structure.
"""
import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from analyze_leaderboard import SOURCE_LABELS
from utils import read_config

# ============================================================
# Style constants
# ============================================================
_config = None
def _get_style():
    global _config
    if _config is None:
        _config = read_config("leaderboard.json")
    style = _config.get("report", {}).get("style", {})
    return style

FONT_CN = "微软雅黑"
FONT_EN = "Arial"
COLOR_HEADER_BG = "1B3A5C"
COLOR_HEADER_TEXT = "FFFFFF"
COLOR_ROW_ALT = "F2F7FB"
COLOR_SECTION = RGBColor(0x1B, 0x3A, 0x5C)
COLOR_ACCENT = RGBColor(0x25, 0x63, 0xEB)
COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)

SRC_COLORS = {
    "lm": RGBColor(0x25, 0x63, 0xEB),
    "aa": RGBColor(0x16, 0xA3, 0x4A),
    "sc": RGBColor(0xDC, 0x26, 0x26),
}


# ============================================================
# Utility functions
# ============================================================
def _set_run(run, font_cn=FONT_CN, font_en=FONT_EN, size=Pt(10),
             bold=False, color=None):
    run.font.name = font_en
    run.font.size = size
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn("w:eastAsia"), font_cn)
    if color:
        run.font.color.rgb = color


def _set_cell(cell, text, bold=False, color=None, size=Pt(9), align="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(str(text))
    _set_run(run, size=size, bold=bold, color=color)


def _shade_cell(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_section(doc, title, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    prefix = "■" if level == 1 else "▸"
    run = p.add_run(f"{prefix} {title}")
    size = Pt(14) if level == 1 else Pt(12)
    _set_run(run, size=size, bold=True, color=COLOR_SECTION)


def _add_text(doc, text, size=Pt(10), color=None, bold=False, space_after=Pt(4)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    _set_run(run, size=size, color=color, bold=bold)
    return p


def _add_image(doc, path, width_inches=6.0):
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))


# ============================================================
# Report generation
# ============================================================
def generate_report(analysis: dict, chart_paths: dict, output_dir: str) -> str:
    date_str = analysis["date"]
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT_EN
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    insights_text = analysis.get("insights", "")

    # Title
    _add_title(doc, date_str)

    # 1. 宏观格局
    _add_section(doc, "宏观格局")
    macro = _extract_section(insights_text, "MACRO_LANDSCAPE")
    if macro:
        _add_text(doc, macro, size=Pt(10))
    else:
        _add_text(doc, "（宏观分析生成中...）", color=COLOR_MUTED)

    # 2. 核心洞察
    _add_section(doc, "核心洞察")
    _add_insights(doc, insights_text)

    # 3. 赛道分析
    _add_section(doc, "赛道分析")
    comparisons = analysis.get("comparisons", {})
    track_charts = chart_paths.get("track_charts", {})

    for track_label, track_data in comparisons.items():
        _add_section(doc, track_label, level=2)
        _add_cross_table(doc, track_data.get("cross_models", []))
        chart_path = track_charts.get(track_label)
        _add_image(doc, chart_path, width_inches=6.5)

        track_short = track_label.split("(")[0].strip()
        track_insight = _extract_track_insight(insights_text, track_short)
        if track_insight:
            _add_text(doc, track_insight, size=Pt(10), color=COLOR_MUTED)

    # 4. 跨平台排名一致性
    _add_section(doc, "跨平台排名一致性")
    _add_text(doc, "散点图展示同一模型在不同排行榜的排名位置。靠近对角线表示跨平台排名一致，偏离对角线表示存在分歧。",
              size=Pt(9), color=COLOR_MUTED)
    _add_image(doc, chart_paths.get("scatter_chart"), width_inches=5.5)

    # 5. 厂商全景
    _add_section(doc, "厂商入榜实力全景")
    _add_vendor_table(doc, analysis.get("vendors", {}))
    _add_image(doc, chart_paths.get("vendor_chart"), width_inches=6.0)

    # 6. 技术壁垒
    _add_section(doc, "技术壁垒")
    tech = analysis.get("tech_barriers", {})
    if tech:
        _add_tech_barriers_summary(doc, tech)
    tech_insight = _extract_section(insights_text, "TECH_BARRIERS")
    if tech_insight:
        _add_text(doc, tech_insight, size=Pt(10))

    # 7. 独有赛道
    exclusives = analysis.get("exclusives", {})
    if exclusives:
        _add_section(doc, "独有赛道摘要")
        excl_insight = _extract_exclusive_insight(insights_text)
        if excl_insight:
            _add_text(doc, excl_insight, size=Pt(10))
        for track_label, info in exclusives.items():
            _add_section(doc, f"{track_label}（{info['source']}）", level=2)
            _add_exclusive_table(doc, info.get("top10", []), info.get("source", ""))

    # 8. 机会筛选
    _add_section(doc, "机会筛选")
    opps = analysis.get("opportunities", {})
    if opps:
        _add_opportunity_table(doc, opps)
    opp_insight = _extract_section(insights_text, "OPPORTUNITY")
    if opp_insight:
        _add_text(doc, opp_insight, size=Pt(10))

    # 9. 结论
    _add_section(doc, "结论")
    conclusion = _extract_section(insights_text, "CONCLUSION")
    if conclusion:
        for line in conclusion.split("\n"):
            line = line.strip()
            if line:
                _add_text(doc, line, size=Pt(11), bold=True)
    else:
        _add_text(doc, "（结论生成中...）", color=COLOR_MUTED)

    # Data sources
    _add_section(doc, "数据来源")
    _add_text(doc, "本报告数据来源于以下三个排行榜平台：", size=Pt(9))
    sources_info = [
        ("LMArena (arena.ai)", "全球最大的 AI 模型匿名对战投票平台，Elo 排名"),
        ("ArtificialAnalysis.ai", "专注 AI 模型基准测试与定价分析，覆盖文本/图像/视频"),
        ("SuperCLUE (superclueai.com)", "中文 AI 模型评测基准，覆盖多模态生成与编辑"),
    ]
    for name, desc in sources_info:
        p = doc.add_paragraph()
        run1 = p.add_run(f"• {name}: ")
        _set_run(run1, size=Pt(9), bold=True)
        run2 = p.add_run(desc)
        _set_run(run2, size=Pt(9), color=COLOR_MUTED)

    # Save
    filename = f"leaderboard_report_{date_str}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f"\n[OK] 报告生成: {filepath}")
    return filepath


# ============================================================
# Internal rendering functions
# ============================================================
def _add_title(doc, date_str):
    try:
        dt = datetime.strptime(date_str, "%Y%m%d")
        date_display = dt.strftime("%Y年%m月%d日")
        week_display = f"W{dt.isocalendar()[1]}"
    except ValueError:
        date_display = date_str
        week_display = ""

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI 模型竞技格局报告")
    _set_run(run, size=Pt(22), bold=True, color=COLOR_SECTION)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"{date_display} | {week_display} | 三源综合分析")
    _set_run(run2, size=Pt(11), color=COLOR_MUTED)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("数据来源: LMArena · ArtificialAnalysis · SuperCLUE")
    _set_run(run3, size=Pt(9), color=COLOR_ACCENT)

    doc.add_paragraph("─" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_insights(doc, insights_text):
    if not insights_text:
        _add_text(doc, "（洞察生成不可用）")
        return

    insights = re.findall(
        r'\[INSIGHT_\d+\]\s*\n*标题[：:]\s*(.*?)\n+正文[：:]\s*(.*?)(?=\[INSIGHT_|\[TRACK_|\[EXCLUSIVE_|\Z)',
        insights_text, re.DOTALL
    )

    if insights:
        for i, (title, body) in enumerate(insights):
            title = title.strip()
            body = body.strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            run_num = p.add_run(f"{i+1}. ")
            _set_run(run_num, size=Pt(11), bold=True, color=COLOR_ACCENT)
            run_title = p.add_run(title)
            _set_run(run_title, size=Pt(11), bold=True)
            for line in body.split("\n"):
                line = line.strip()
                if line:
                    _add_text(doc, line, size=Pt(10))
    else:
        for line in insights_text.split("\n"):
            line = line.strip()
            if line and not line.startswith("["):
                _add_text(doc, line, size=Pt(10))


def _extract_track_insight(insights_text: str, track_short: str) -> str:
    if not insights_text:
        return ""
    pattern = rf'\[TRACK_{re.escape(track_short)}\]\s*\n*(.*?)(?=\[TRACK_|\[EXCLUSIVE_|\[INSIGHT_|\[TECH_|\[OPPORTUNITY|\[CONCLUSION|\[MACRO_|\Z)'
    match = re.search(pattern, insights_text, re.DOTALL)
    if match:
        text = match.group(1).strip()
        lines = [l.strip() for l in text.split("\n") if l.strip() and not l.strip().startswith("[")]
        return " ".join(lines)
    return ""


def _extract_exclusive_insight(insights_text: str) -> str:
    if not insights_text:
        return ""
    pattern = r'\[EXCLUSIVE_TRACKS\]\s*\n*(.*?)(?=\[TRACK_|\[INSIGHT_|\[TECH_|\[OPPORTUNITY|\[CONCLUSION|\[MACRO_|\Z)'
    match = re.search(pattern, insights_text, re.DOTALL)
    if match:
        text = match.group(1).strip()
        lines = [l.strip() for l in text.split("\n") if l.strip() and not l.strip().startswith("[")]
        return " ".join(lines)
    return ""


def _extract_section(insights_text: str, marker: str) -> str:
    if not insights_text:
        return ""
    pattern = rf'\[{re.escape(marker)}\]\s*\n*(.*?)(?=\[(?:INSIGHT_|TRACK_|EXCLUSIVE_|TECH_|OPPORTUNITY|CONCLUSION|MACRO_)|\Z)'
    match = re.search(pattern, insights_text, re.DOTALL)
    if match:
        text = match.group(1).strip()
        lines = [l.strip() for l in text.split("\n") if l.strip() and not l.strip().startswith("[") and not l.strip().startswith("===")]
        return " ".join(lines)
    return ""


def _add_tech_barriers_summary(doc, tech):
    open_c = tech.get("open_count", 0)
    closed_c = tech.get("closed_count", 0)
    total = open_c + closed_c
    if total == 0:
        return
    open_pct = round(open_c / total * 100)
    _add_text(doc,
              f"Top 20 模型中：开源 {open_c} 个（{open_pct}%）| 闭源 {closed_c} 个（{100-open_pct}%）"
              f" | 中国厂商开源 {tech.get('chinese_open', 0)} 个、闭源 {tech.get('chinese_closed', 0)} 个",
              size=Pt(9), color=COLOR_MUTED)


def _add_opportunity_table(doc, opps):
    opportunities = opps.get("opportunities", [])
    if not opportunities:
        return
    chinese = opps.get("chinese_highlights", [])
    if chinese:
        _add_section(doc, "值得关注的中国厂商", level=2)
        _build_opp_table(doc, chinese)
    global_l = opps.get("global_leaders", [])
    if global_l:
        _add_section(doc, "全球领先厂商", level=2)
        _build_opp_table(doc, global_l)


def _build_opp_table(doc, items):
    headers = ["厂商", "关注度", "标签", "优势赛道"]
    table = doc.add_table(rows=len(items) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        _set_cell(table.rows[0].cells[i], h, bold=True,
                  color=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(8))
        _shade_cell(table.rows[0].cells[i], COLOR_HEADER_BG)

    for idx, item in enumerate(items):
        row = table.rows[idx + 1]
        if idx % 2 == 1:
            for cell in row.cells:
                _shade_cell(cell, COLOR_ROW_ALT)
        _set_cell(row.cells[0], item["vendor"], size=Pt(9), bold=True, align="left")
        stars = "★" * min(item["score"], 5) + "☆" * max(0, 5 - item["score"])
        _set_cell(row.cells[1], stars, size=Pt(9))
        tags_str = " ".join(item.get("tags", [])[:2])
        _set_cell(row.cells[2], tags_str, size=Pt(7), align="left")
        tracks_str = ", ".join(item.get("top_tracks", [])[:2])
        _set_cell(row.cells[3], tracks_str, size=Pt(7), align="left")


def _add_cross_table(doc, cross_models):
    if not cross_models:
        _add_text(doc, "（暂无跨源对比数据）")
        return

    headers = ["Model", "Vendor", "LMArena", "AA", "SuperCLUE"]
    table = doc.add_table(rows=min(len(cross_models), 15) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        _set_cell(table.rows[0].cells[i], h, bold=True,
                  color=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(9))
        _shade_cell(table.rows[0].cells[i], COLOR_HEADER_BG)

    for idx, cm in enumerate(cross_models[:15]):
        row = table.rows[idx + 1]
        if idx % 2 == 1:
            for cell in row.cells:
                _shade_cell(cell, COLOR_ROW_ALT)

        model = cm.get("model", "?")
        if len(model) > 30:
            model = model[:28] + "…"
        _set_cell(row.cells[0], model, size=Pt(8), align="left")
        _set_cell(row.cells[1], cm.get("vendor", "?"), size=Pt(8))

        for col_idx, src in enumerate(["lm", "aa", "sc"], start=2):
            rank = cm.get(f"{src}_rank")
            score = cm.get(f"{src}_score")
            if rank is not None:
                text = f"#{rank} ({score})"
                _set_cell(row.cells[col_idx], text, size=Pt(8),
                          color=SRC_COLORS.get(src))
            else:
                _set_cell(row.cells[col_idx], "—", size=Pt(8),
                          color=COLOR_MUTED)


def _add_vendor_table(doc, vendor_data):
    if not vendor_data:
        return
    top_vendors = list(vendor_data.items())[:10]
    headers = ["厂商", "入榜数", "覆盖赛道", "最佳排名"]
    table = doc.add_table(rows=len(top_vendors) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        _set_cell(table.rows[0].cells[i], h, bold=True,
                  color=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(9))
        _shade_cell(table.rows[0].cells[i], COLOR_HEADER_BG)

    for idx, (vendor, info) in enumerate(top_vendors):
        row = table.rows[idx + 1]
        if idx % 2 == 1:
            for cell in row.cells:
                _shade_cell(cell, COLOR_ROW_ALT)
        _set_cell(row.cells[0], vendor, size=Pt(9), bold=True, align="left")
        _set_cell(row.cells[1], str(info["total_entries"]), size=Pt(9))
        tracks_str = ", ".join(
            f"{t.split('(')[0].strip()}:#{d['best_rank']}"
            for t, d in list(info["tracks"].items())[:4]
        )
        if len(info["tracks"]) > 4:
            tracks_str += f" +{len(info['tracks'])-4}"
        _set_cell(row.cells[2], tracks_str, size=Pt(7), align="left")
        best = info["best_overall_rank"]
        _set_cell(row.cells[3], f"#{best}" if best < 999 else "—", size=Pt(9))


def _add_exclusive_table(doc, top10, source_name):
    if not top10:
        return
    score_key = "median" if "SuperCLUE" in source_name else "score"
    table = doc.add_table(rows=min(len(top10), 10) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Rank", "Model", "Score", "Org/Vendor"]
    for i, h in enumerate(headers):
        _set_cell(table.rows[0].cells[i], h, bold=True,
                  color=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(8))
        _shade_cell(table.rows[0].cells[i], COLOR_HEADER_BG)

    for idx, row_data in enumerate(top10[:10]):
        row = table.rows[idx + 1]
        if idx % 2 == 1:
            for cell in row.cells:
                _shade_cell(cell, COLOR_ROW_ALT)
        _set_cell(row.cells[0], row_data.get("rank", "?"), size=Pt(8))
        model = row_data.get("model", "?")
        if len(model) > 25:
            model = model[:23] + "…"
        _set_cell(row.cells[1], model, size=Pt(8), align="left")
        _set_cell(row.cells[2], row_data.get(score_key, row_data.get("score", "?")), size=Pt(8))
        _set_cell(row.cells[3], row_data.get("org", row_data.get("creator", "?")), size=Pt(8))
