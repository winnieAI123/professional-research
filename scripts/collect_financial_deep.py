"""
Type 8: Financial Data Extraction (财务数据提取)
Dynamic multi-company financial data collection pipeline.

Usage:
  python collect_financial_deep.py --companies "蚂蚁集团,微众银行,奇富科技" \
    --query "分产品余额、收入和利润" --years 5 --output results/report

No hardcoded company config — auto-detects data source for each company.
"""

import argparse
import json
import os
import re
import sys
import time
import requests
from datetime import datetime
from bs4 import BeautifulSoup

# Add scripts dir to path for sibling imports
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SCRIPT_DIR)

from llm_client import generate_content, get_client, FAST_MODEL

# Fast model fallback chain (for table filtering / classification)
FAST_MODEL_CHAIN = [
    "models/gemini-3-flash-preview",
    "models/gemini-2.5-flash",
]
from utils import get_api_key

# ============================================================
# Constants
# ============================================================
SEC_USER_AGENT = "CompanyResearch research@example.com"
SEC_RATE_LIMIT = 0.15  # seconds between SEC requests
TAVILY_API_KEY = get_api_key("TAVILY_API_KEY")


# ============================================================
# JSON Normalization Helpers
# ============================================================
def _normalize_extracted_json(parsed: dict) -> dict:
    """
    Normalize LLM output to standard format: {"data": {...}, "metadata": {...}}.
    Handles cases where LLM uses custom top-level keys or array-based data.
    """
    if not parsed or not isinstance(parsed, dict):
        return parsed
    
    # Case 1: Already has "data" key with dict value
    if "data" in parsed and isinstance(parsed["data"], dict):
        return parsed
    
    # Case 2: LLM used a custom key (e.g., "蚂蚁集团财务数据")
    # Find the first key that looks like a data container (not "metadata")
    metadata = parsed.get("metadata", {})
    data_candidates = {}
    
    for key, value in parsed.items():
        if key == "metadata":
            continue
        if isinstance(value, dict):
            # Check if this dict contains year-keyed data or nested metrics
            has_metrics = False
            normalized_group = {}
            for sub_key, sub_val in value.items():
                if isinstance(sub_val, dict):
                    # Check if it's {year: value} format
                    if any(str(y).isdigit() and len(str(y)) == 4 for y in sub_val.keys()):
                        has_metrics = True
                        normalized_group[sub_key] = sub_val
                    else:
                        # Nested group — recurse one level
                        inner_metrics = {}
                        for inner_k, inner_v in sub_val.items():
                            if isinstance(inner_v, dict):
                                inner_metrics[inner_k] = inner_v
                            elif isinstance(inner_v, list):
                                # Convert array format to {year: value} dict
                                converted = _array_to_year_dict(inner_v)
                                if converted:
                                    inner_metrics[inner_k] = converted
                            elif isinstance(inner_v, (int, float, str)):
                                # Skip metadata-like fields (unit, notes, etc.)
                                pass
                        if inner_metrics:
                            has_metrics = True
                            normalized_group[sub_key] = inner_metrics
                elif isinstance(sub_val, list):
                    converted = _array_to_year_dict(sub_val)
                    if converted:
                        has_metrics = True
                        normalized_group[sub_key] = converted
                elif isinstance(sub_val, str):
                    # Data not found message — skip
                    pass
            
            if has_metrics:
                data_candidates[key] = normalized_group
    
    if data_candidates:
        return {"data": data_candidates, "metadata": metadata}
    
    return parsed


def _array_to_year_dict(arr: list) -> dict:
    """Convert [{年份: 2020, 指标: 值}, ...] array to {2020: 值, ...} dict."""
    if not arr or not isinstance(arr, list):
        return {}
    
    result = {}
    for item in arr:
        if not isinstance(item, dict):
            continue
        # Find year field
        year = None
        value = None
        for k, v in item.items():
            k_lower = k.lower()
            if '年' in k or 'year' in k_lower:
                year = str(v)
            elif '利润' in k or 'profit' in k_lower or '收入' in k_lower or 'revenue' in k_lower or 'income' in k_lower:
                value = v
        
        if year and value is not None:
            result[year] = value
    
    return result


def _try_reformat_response(raw_text: str) -> dict:
    """Ask LLM to reformat a non-JSON response into standard format."""
    try:
        prompt = f"""以下文本包含财务数据但格式不标准。请将其转换为以下JSON格式:
{{
  "data": {{
    "数据组名": {{
      "指标名": {{"2020": 值, "2021": 值}}
    }}
  }},
  "metadata": {{"unit": "单位", "notes": ""}}
}}

原始文本:
{raw_text[:10000]}"""
        
        result = generate_content(prompt=prompt, max_output_tokens=4000, model=FAST_MODEL)
        if result:
            parsed = json.loads(result.strip().strip('```json').strip('```').strip())
            return _normalize_extracted_json(parsed)
    except Exception:
        pass
    return None


# ============================================================
# Dynamic Company Detection (NO hardcoded config)
# ============================================================
def detect_data_source(company_name: str) -> dict:
    """
    Auto-detect best data source for a company.
    Returns: {"name": ..., "source": "sec_edgar"|"cn_listed"|"pdf_search"|"web_search", ...}
    """
    print(f"\n  [Detect] Analyzing data source for '{company_name}'...")
    
    # 1. Try US stock ticker resolution
    us_info = _try_resolve_us_ticker(company_name)
    if us_info:
        print(f"    ✓ US-listed: {us_info['ticker']} → SEC EDGAR")
        return {
            "name": company_name,
            "source": "sec_edgar",
            "ticker": us_info["ticker"],
            "filing_type": us_info.get("filing_type", "20-F"),
        }
    
    # 2. Try CN-listed stock (EastMoney search)
    cn_ticker = _try_resolve_cn_ticker(company_name)
    if cn_ticker:
        print(f"    ✓ CN-listed: {cn_ticker} → EastMoney F10")
        return {
            "name": company_name,
            "source": "cn_listed",
            "ticker": cn_ticker,
        }
    
    # 3. Ask LLM: is this a bank/financial institution with public annual reports?
    if _is_likely_public_reporter(company_name):
        print(f"    → Non-listed institution with annual reports → PDF search")
        return {
            "name": company_name,
            "source": "pdf_search",
            "search_name": company_name,
        }
    
    # 4. Fallback: web search
    print(f"    → No public filings detected → Web search")
    return {
        "name": company_name,
        "source": "web_search",
        "search_name": company_name,
    }


def _try_resolve_us_ticker(company_name: str) -> dict:
    """Check if company is US-listed via yfinance."""
    try:
        import yfinance as yf
        # If input looks like a ticker already (all uppercase alpha)
        if re.match(r'^[A-Z]{1,5}$', company_name):
            t = yf.Ticker(company_name)
            info = t.info
            if info.get("regularMarketPrice"):
                country = info.get("country", "")
                filing = "10-K" if country == "United States" else "20-F"
                return {"ticker": company_name, "filing_type": filing}
        
        # Try common Chinese fintech → US ticker mappings via web search
        prompt = f"""Given the company name "{company_name}", if it is listed on a US stock exchange (NYSE/NASDAQ), 
return ONLY its ticker symbol (e.g., "QFIN"). If it is NOT US-listed, return "NONE".
Just the ticker or NONE, nothing else."""
        result = generate_content(prompt=prompt, max_output_tokens=20, model=FAST_MODEL)
        if result and result.strip() != "NONE" and re.match(r'^[A-Z]{1,5}$', result.strip()):
            ticker = result.strip()
            t = yf.Ticker(ticker)
            info = t.info
            if info.get("regularMarketPrice"):
                country = info.get("country", "")
                filing = "10-K" if country == "United States" else "20-F"
                return {"ticker": ticker, "filing_type": filing}
    except Exception as e:
        print(f"    US ticker check failed: {e}")
    return None


def _try_resolve_cn_ticker(company_name: str) -> str:
    """Check if company is CN-listed via EastMoney search."""
    try:
        url = "https://searchapi.eastmoney.com/api/suggest/get"
        params = {
            "input": company_name,
            "type": "14",
            "token": "D43BF722C8E33BDC906FB84D85E326E8",
            "count": "5",
        }
        resp = requests.get(url, params=params, timeout=10)
        data = resp.json()
        if data.get("QuotationCodeTable", {}).get("Data"):
            for item in data["QuotationCodeTable"]["Data"]:
                code = item.get("Code", "")
                market = item.get("MktNum", "")
                if market in ("0", "1") and re.match(r'^\d{6}$', code):
                    return code
    except Exception:
        pass
    return None


def _is_likely_public_reporter(company_name: str) -> bool:
    """Ask LLM if this company likely publishes annual reports (banks, etc)."""
    try:
        prompt = f"""Is "{company_name}" a licensed financial institution (bank, consumer finance company, 
insurance company) in China that is required to publish annual reports publicly?
Answer ONLY "YES" or "NO"."""
        result = generate_content(prompt=prompt, max_output_tokens=10, model=FAST_MODEL)
        return result and "YES" in result.upper()
    except Exception:
        return False


# ============================================================
# Path A: SEC EDGAR (US-listed)
# ============================================================
class SECCollector:
    """Lightweight SEC EDGAR collector."""
    
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': SEC_USER_AGENT,
            'Accept-Encoding': 'gzip, deflate',
        })
        self._last_request = 0
    
    def _rate_limit(self):
        elapsed = time.time() - self._last_request
        if elapsed < SEC_RATE_LIMIT:
            time.sleep(SEC_RATE_LIMIT - elapsed)
        self._last_request = time.time()
    
    def _get(self, url):
        for attempt in range(3):
            self._rate_limit()
            try:
                resp = self.session.get(url, timeout=30)
                resp.raise_for_status()
                return resp.json()
            except requests.exceptions.HTTPError as e:
                if resp.status_code in (503, 429) and attempt < 2:
                    wait = [5, 15, 30][attempt]
                    print(f"      [SEC] {resp.status_code} on attempt {attempt+1}, retrying in {wait}s...")
                    time.sleep(wait)
                else:
                    raise
    
    def lookup_cik(self, ticker: str) -> str:
        data = self._get('https://www.sec.gov/files/company_tickers.json')
        for entry in data.values():
            if entry.get('ticker', '').upper() == ticker.upper():
                return str(entry['cik_str']).zfill(10)
        return None
    
    def get_filings(self, cik: str, filing_type: str = '20-F') -> list:
        data = self._get(f"https://data.sec.gov/submissions/CIK{cik}.json")
        filings = data['filings']['recent']
        results = []
        for i, form in enumerate(filings['form']):
            if form == filing_type:
                results.append({
                    'date': filings['filingDate'][i],
                    'accession': filings['accessionNumber'][i],
                    'primary_doc': filings['primaryDocument'][i],
                    'size': filings.get('size', [0] * len(filings['form']))[i],
                })
        return results
    
    def download_filing(self, cik: str, filing: dict, output_dir: str) -> str:
        accession = filing['accession'].replace('-', '')
        filename = f"{filing['date']}_{filing['primary_doc']}"
        filepath = os.path.join(output_dir, filename)
        if os.path.exists(filepath) and os.path.getsize(filepath) > 1000:
            return filepath
        
        url = f"https://www.sec.gov/Archives/edgar/data/{cik}/{accession}/{filing['primary_doc']}"
        for attempt in range(3):
            self._rate_limit()
            try:
                resp = self.session.get(url, timeout=60)
                resp.raise_for_status()
                os.makedirs(output_dir, exist_ok=True)
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(resp.text)
                return filepath
            except requests.exceptions.HTTPError as e:
                if resp.status_code in (503, 429) and attempt < 2:
                    wait = [5, 15, 30][attempt]
                    print(f"      [SEC] {resp.status_code} downloading {filename}, retry in {wait}s...")
                    time.sleep(wait)
                else:
                    raise


    def get_filing_documents(self, cik: str, accession: str) -> list:
        """Get all documents in a filing by parsing EDGAR HTML directory.
        Returns list of dicts with keys: name, type.
        Uses cik_int (no leading zeros) for archive URLs.
        """
        import re as _re
        cik_int = str(int(cik))  # Strip leading zeros for archive URLs
        acc_nodash = accession.replace('-', '')
        url = f"https://www.sec.gov/Archives/edgar/data/{cik_int}/{acc_nodash}/"
        try:
            self._rate_limit()
            resp = self.session.get(url, timeout=30)
            resp.raise_for_status()
            # Parse all archive file links from the HTML directory listing
            links = _re.findall(
                r'href="/Archives/edgar/data/[^"]+/([^"/]+\.(?:htm|html|xml|txt))"',
                resp.text, _re.IGNORECASE
            )
            # Also find exhibit type labels (table cells near file links)
            # Pattern: files with ex99 or ex-99 in name are EX-99.1
            items = []
            for fname in links:
                ftype = ''
                fname_lower = fname.lower()
                if 'ex99-1' in fname_lower or 'ex99_1' in fname_lower or 'ex-99-1' in fname_lower:
                    ftype = 'EX-99.1'
                elif 'ex99' in fname_lower:
                    ftype = 'EX-99'
                elif '_6k' in fname_lower or fname_lower.endswith('_6k.htm'):
                    ftype = '6-K'
                items.append({'name': fname, 'type': ftype})
            return items
        except Exception:
            return []

    def download_exhibit_file(self, cik: str, accession: str, exhibit_type: str,
                              filing_date: str, output_dir: str) -> str:
        """Download a specific exhibit from a 6-K filing.

        For foreign private issuers, the earnings press release is typically
        filed as exhibit EX-99.1 and contains the full financial tables.

        exhibit_type: 'EX-99.1' for earnings press release
        Returns file path if successful, None otherwise.
        """
        acc_nodash = accession.replace('-', '')
        items = self.get_filing_documents(cik, accession)
        for item in items:
            item_type = item.get('type', '').upper().replace(' ', '').replace('-', '')
            target_type = exhibit_type.upper().replace(' ', '').replace('-', '')
            if item_type == target_type:
                filename = item.get('name', '')
                if not filename:
                    continue
                filepath = os.path.join(output_dir, f"{filing_date}_{filename}")
                if os.path.exists(filepath) and os.path.getsize(filepath) > 5000:
                    return filepath
                url = f"https://www.sec.gov/Archives/edgar/data/{str(int(cik))}/{acc_nodash}/{filename}"
                try:
                    self._rate_limit()
                    resp = self.session.get(url, timeout=60)
                    resp.raise_for_status()
                    os.makedirs(output_dir, exist_ok=True)
                    with open(filepath, 'w', encoding='utf-8') as f:
                        f.write(resp.text)
                    return filepath
                except Exception as e:
                    print(f"      [SEC] Failed to download exhibit {exhibit_type}: {e}")
                    return None
        return None  # exhibit not found in this filing


def collect_sec_edgar(company: dict, query: str, target_years: list, data_dir: str) -> dict:
    """SEC EDGAR pipeline: download 20-F/10-K → LLM two-round extract."""
    ticker = company["ticker"]
    filing_type = company.get("filing_type", "20-F")
    
    collector = SECCollector()
    cik = collector.lookup_cik(ticker)
    if not cik:
        print(f"    ✗ Could not find CIK for {ticker}")
        return None
    
    filings = collector.get_filings(cik, filing_type)
    if not filings:
        print(f"    ✗ No {filing_type} filings found")
        return None
    
    print(f"    Found {len(filings)} {filing_type} filings")
    
    raw_dir = os.path.join(data_dir, "raw", ticker.lower())
    os.makedirs(raw_dir, exist_ok=True)
    
    # Download enough filings to cover target years
    download_count = min((len(target_years) + 2) // 3 + 1, len(filings))
    html_paths = []
    for f in filings[:download_count]:
        path = collector.download_filing(cik, f, raw_dir)
        if path:
            html_paths.append(path)
            print(f"    ✓ Downloaded: {os.path.basename(path)}")
    
    if not html_paths:
        return None
    
    # Two-round LLM extraction
    return _extract_from_html_files(html_paths, query, target_years, filings, collector, cik, raw_dir)


def collect_sec_quarterly(company: dict, query: str, num_quarters: int, data_dir: str) -> dict:
    """
    SEC EDGAR quarterly pipeline: download recent 10-Q or 6-K filings → extract.
    
    For US domestic companies: 10-Q (standard quarterly report)
    For foreign private issuers (e.g. Chinese ADRs): 6-K (interim reports)
    
    6-K filings contain various announcements, not all financial.
    Strategy: download → check table count → only extract if financial.
    """
    ticker = company["ticker"]
    annual_filing_type = company.get("filing_type", "20-F")
    
    # Determine quarterly filing type based on annual type
    if annual_filing_type == "10-K":
        quarterly_type = "10-Q"
    else:
        quarterly_type = "6-K"  # Foreign private issuers
    
    collector = SECCollector()
    cik = collector.lookup_cik(ticker)
    if not cik:
        print(f"    ✗ Could not find CIK for {ticker}")
        return None
    
    filings = collector.get_filings(cik, quarterly_type)
    if not filings:
        print(f"    ✗ No {quarterly_type} filings found")
        return None
    
    print(f"    Found {len(filings)} {quarterly_type} filings, processing recent {num_quarters}...")
    
    raw_dir = os.path.join(data_dir, "raw", f"{ticker.lower()}_quarterly")
    os.makedirs(raw_dir, exist_ok=True)
    
    quarterly_data = {}
    quarterly_metadata = {"source": f"sec_edgar:{quarterly_type}", "period": "quarterly"}
    extracted_count = 0
    
    # For 10-Q: straightforward, download and extract
    # For 6-K: foreign private issuers file many 6-Ks (not all financial).
    #   Strategy: for each 6-K, first try exhibit 99.1 (earnings press release,
    #   ~100KB+ with full tables). Fall back to primary doc if no exhibit.
    #   Expand pool to num_quarters*10 to find enough earnings 6-Ks.
    candidates = filings[:num_quarters * 10] if quarterly_type == "6-K" else filings[:num_quarters]

    financial_keywords = ['revenue', 'income', 'profit', 'loss', 'balance',
                          'assets', 'liabilities', 'cash flow', 'earnings',
                          '\u6536\u5165', '\u5229\u6da6', '\u8d44\u4ea7', '\u8d1f\u503a']

    # For 6-K: pre-filter by submission size.
    # Earnings filings (cover + press release) are typically >200KB total.
    # Cover-only 6-Ks are <20KB. This avoids fetching index for every filing.
    if quarterly_type == "6-K":
        candidates = [f for f in candidates
                      if not isinstance(f.get("size", 0), (int, float))
                      or f.get("size", 0) > 100000]
        if not candidates:
            candidates = filings[:num_quarters * 10]  # fallback if no size info
        print(f"    After size filter: {len(candidates)} candidate 6-Ks")

    for filing in candidates:
        if extracted_count >= num_quarters:
            break

        path = None
        if quarterly_type == "6-K":
            # Try exhibit 99.1 first (earnings press release with financials)
            path = collector.download_exhibit_file(
                cik, filing['accession'], 'EX-99.1', filing['date'], raw_dir)
            if path:
                print(f"    [Q] Found exhibit 99.1: {os.path.basename(path)} ({filing['date']})")
            else:
                # Fall back to primary doc; skip if it's just a small cover page
                primary = collector.download_filing(cik, filing, raw_dir)
                if not primary:
                    continue
                if os.path.getsize(primary) < 15000:
                    continue  # Cover-page 6-Ks are <15KB, skip
                path = primary

            # Validate: must have financial tables
            tables = _html_to_tables(path)
            has_financial = False
            if len(tables) >= 3:
                sample_text = ' '.join(t['markdown'][:300].lower() for t in tables[:10])
                if any(kw in sample_text for kw in financial_keywords):
                    has_financial = True
            if not has_financial:
                continue  # Not a financial filing, skip
        else:
            path = collector.download_filing(cik, filing, raw_dir)
            if not path:
                continue

        print(f"    [Q] Extracting: {os.path.basename(path)} ({filing['date']})")
        
        # Use same extraction pipeline (filter + extract)
        tables = _html_to_tables(path)
        if not tables:
            continue
        
        relevant = _filter_tables(tables, query)
        if not relevant:
            continue
        
        # Extract with quarterly-aware prompt
        result = _extract_quarterly_data(relevant, query, filing['date'])
        if result and result.get("data"):
            for group, metrics in result["data"].items():
                q_group = f"[季度] {group}" if not group.startswith("[季度]") else group
                if q_group not in quarterly_data:
                    quarterly_data[q_group] = {}
                for metric, periods in metrics.items():
                    if metric not in quarterly_data[q_group]:
                        quarterly_data[q_group][metric] = {}
                    quarterly_data[q_group][metric].update(periods)
            
            extracted_count += 1
            print(f"      ✓ Extracted quarterly data ({extracted_count}/{num_quarters})")
    
    if not quarterly_data:
        print(f"    ⚠ No quarterly data extracted from {quarterly_type} filings")
        return None
    
    print(f"    ✓ Quarterly total: {extracted_count} filings, {len(quarterly_data)} groups")
    return {"data": quarterly_data, "metadata": quarterly_metadata}


def _extract_quarterly_data(tables: list, query: str, filing_date: str) -> dict:
    """Extract structured data from quarterly filings (10-Q / 6-K)."""
    tables_text = ""
    for t in tables:
        tables_text += f"\n--- Table #{t['index']} (context: {t['context']}) ---\n{t['markdown']}\n"
    
    if len(tables_text) > 100000:
        tables_text = tables_text[:100000]
    
    prompt = f"""从以下季度财务报表表格中提取结构化数据。

用户需求: {query}
报告日期: {filing_date}

这是季度报告（10-Q 或 6-K），请注意：
1. 提取各季度的数据
2. 期间的key用"YYYY-QN"格式（如"2025-Q1", "2024-Q3"）
3. 如果表格同时有季度和年累计数据，优先提取季度数据
4. 找不到的不编造

输出JSON格式:
{{
  "data": {{
    "数据组名": {{
      "指标名(English)": {{"2025-Q1": 值, "2024-Q4": 值, "2024-Q3": 值}}
    }}
  }},
  "metadata": {{"unit": "单位", "filing_date": "{filing_date}"}}
}}

表格内容:
{tables_text}"""
    
    result = generate_content(prompt=prompt, max_output_tokens=8000)
    if not result:
        return None
    
    try:
        cleaned = result.strip().strip('```json').strip('```').strip()
        parsed = json.loads(cleaned)
        return _normalize_extracted_json(parsed)
    except Exception as e:
        print(f"      ✗ Quarterly JSON parse failed: {e}")
        return None


def _extract_from_html_files(html_paths, query, target_years, all_filings=None,
                              collector=None, cik=None, raw_dir=None) -> dict:
    """Two-round LLM extraction from SEC HTML filings."""
    all_data = {}
    all_metadata = {}
    
    for html_path in html_paths:
        print(f"    [Extract] {os.path.basename(html_path)}...")
        
        # Round 1: Parse HTML tables
        tables = _html_to_tables(html_path)
        if not tables:
            continue
        print(f"      Found {len(tables)} tables")
        
        # Round 1.5: Filter relevant tables (Flash model, fast)
        relevant = _filter_tables(tables, query)
        print(f"      Filtered to {len(relevant)} relevant tables")
        
        if not relevant:
            continue
        
        # Round 2: Extract structured data (Pro model, precise)
        result = _extract_structured_data(relevant, query, target_years)
        if result and result.get("data"):
            for group, metrics in result["data"].items():
                if group not in all_data:
                    all_data[group] = {}
                for metric, years in metrics.items():
                    if metric not in all_data[group]:
                        all_data[group][metric] = {}
                    all_data[group][metric].update(years)
            if result.get("metadata"):
                all_metadata.update(result["metadata"])
    
    # Check year coverage, download more if needed
    found_years = set()
    for group in all_data.values():
        for metric in group.values():
            found_years.update(str(y) for y in metric.keys())
    
    missing = [str(y) for y in target_years if str(y) not in found_years]
    if missing and all_filings and collector and cik and raw_dir:
        print(f"    Missing years: {missing}, trying older filings...")
        for f in all_filings[len(html_paths):len(html_paths)+2]:
            path = collector.download_filing(cik, f, raw_dir)
            if path:
                extra = _extract_from_html_files([path], query, target_years)
                if extra and extra.get("data"):
                    for g, metrics in extra["data"].items():
                        if g not in all_data:
                            all_data[g] = {}
                        for m, years in metrics.items():
                            if m not in all_data[g]:
                                all_data[g][m] = {}
                            for y, v in years.items():
                                if y not in all_data[g][m]:
                                    all_data[g][m][y] = v
    
    return {"data": all_data, "metadata": all_metadata} if all_data else None


def _html_to_tables(html_path: str) -> list:
    """Parse all <table> from HTML, convert to Markdown format."""
    with open(html_path, 'r', encoding='utf-8', errors='ignore') as f:
        html = f.read()
    
    soup = BeautifulSoup(html, 'html.parser')
    tables = []
    
    for i, table_elem in enumerate(soup.find_all('table')):
        rows = []
        for tr in table_elem.find_all('tr'):
            cells = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
            if any(cells):
                rows.append('| ' + ' | '.join(cells) + ' |')
        
        if len(rows) < 2:
            continue
        
        # Get context (text before table)
        context = ""
        prev = table_elem.find_previous(['h1', 'h2', 'h3', 'h4', 'p', 'b', 'strong'])
        if prev:
            context = prev.get_text(strip=True)[:200]
        
        markdown = '\n'.join(rows)
        tables.append({
            "index": i,
            "context": context,
            "markdown": markdown,
            "preview": '\n'.join(rows[:3]),
            "rows": len(rows),
            "cols": max((r.count('|') - 1 for r in rows), default=0),
        })
    
    return tables


def _filter_tables(tables: list, query: str) -> list:
    """Round 1: Use Flash model chain to quickly filter relevant tables."""
    if len(tables) <= 5:
        return tables
    
    summaries = []
    for t in tables:
        summaries.append(f"Table #{t['index']} ({t['rows']}×{t['cols']}): {t['context'][:100]} | {t['preview'][:150]}")
    
    prompt = f"""User needs: {query}

Below are {len(tables)} tables from a financial filing. Return ONLY the table numbers (e.g., [54, 161, 186]) 
that contain data relevant to the user's query. Return a JSON array of numbers.

{chr(10).join(summaries)}"""

    # Try each fast model in chain (fallback on 429/503)
    for model in FAST_MODEL_CHAIN:
        try:
            result = generate_content(prompt=prompt, max_output_tokens=500, model=model)
            if result:
                indices = json.loads(result.strip().strip('```json').strip('```'))
                return [t for t in tables if t["index"] in indices]
        except Exception:
            continue
    
    # Final fallback: keyword-based filtering
    keywords = ['revenue', 'income', 'profit', 'loan', 'balance', '收入', '利润', '余额']
    return [t for t in tables if any(k in t['markdown'].lower() for k in keywords)][:15]


def _extract_structured_data(tables: list, query: str, target_years: list) -> dict:
    """Round 2: Use Pro model to extract structured JSON data.
    Splits into batches if total table text exceeds 30K chars to avoid
    output truncation (89K prompt → model runs out of output window).
    """
    BATCH_CHAR_LIMIT = 25000  # Keep prompt under 30K to leave room for output
    
    # Build batches
    batches = []
    current_batch = []
    current_chars = 0
    
    for t in tables:
        t_len = len(t['markdown']) + len(t.get('context', '')) + 50
        if current_chars + t_len > BATCH_CHAR_LIMIT and current_batch:
            batches.append(current_batch)
            current_batch = []
            current_chars = 0
        current_batch.append(t)
        current_chars += t_len
    if current_batch:
        batches.append(current_batch)
    
    if len(batches) > 1:
        print(f"      Splitting {len(tables)} tables into {len(batches)} batches")
    
    all_data = {}
    all_metadata = {}
    
    for batch_idx, batch_tables in enumerate(batches):
        tables_text = ""
        for t in batch_tables:
            tables_text += f"\n--- Table #{t['index']} (context: {t['context']}) ---\n{t['markdown']}\n"
        
        batch_label = f"[Batch {batch_idx+1}/{len(batches)}] " if len(batches) > 1 else ""
        
        prompt = f"""从以下财务报表表格中提取结构化数据。

用户需求: {query}
目标年份: {target_years}

规则:
1. 输出严格的JSON格式
2. 指标名使用"中文(English)"格式，如"贷款撮合收入(Revenue from loan facilitation)"
3. 数值保持原始数字，不要转换单位
4. 找不到的年份不要编造
5. 在metadata中注明单位

输出格式:
{{
  "data": {{
    "数据组名": {{
      "指标名1": {{"2020": 123, "2021": 456}},
      "指标名2": {{"2020": 789}}
    }}
  }},
  "metadata": {{
    "unit": "千元RMB",
    "years_found": [2020, 2021],
    "notes": "..."
  }}
}}

表格内容:
{tables_text}"""

        result = generate_content(prompt=prompt, max_output_tokens=8000)
        if not result:
            print(f"      {batch_label}✗ Empty response")
            continue
        
        parsed = _try_parse_json(result)
        if parsed and parsed.get("data"):
            for group, metrics in parsed["data"].items():
                if group not in all_data:
                    all_data[group] = {}
                for metric, years in metrics.items():
                    if isinstance(years, dict):
                        if metric not in all_data[group]:
                            all_data[group][metric] = {}
                        all_data[group][metric].update(years)
            if parsed.get("metadata"):
                all_metadata.update(parsed["metadata"])
            metrics_count = sum(len(m) for g in parsed["data"].values() for m in g.values() if isinstance(m, dict))
            print(f"      {batch_label}✓ {metrics_count} data points extracted")
        else:
            print(f"      {batch_label}✗ Extraction failed")
    
    return {"data": all_data, "metadata": all_metadata} if all_data else None


def _try_parse_json(raw: str) -> dict:
    """Try to parse JSON from LLM output with repair for common issues."""
    cleaned = raw.strip().strip('```json').strip('```').strip()
    
    # Attempt 1: Direct parse
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass
    
    # Attempt 2: Fix truncated JSON by closing unclosed braces
    try:
        open_braces = cleaned.count('{') - cleaned.count('}')
        open_brackets = cleaned.count('[') - cleaned.count(']')
        if open_braces > 0 or open_brackets > 0:
            # Remove last incomplete line
            last_newline = cleaned.rfind('\n')
            if last_newline > len(cleaned) * 0.5:
                truncated = cleaned[:last_newline]
            else:
                truncated = cleaned
            # Close brackets and braces
            truncated += ']' * max(0, open_brackets) + '}' * max(0, open_braces)
            parsed = json.loads(truncated)
            print(f"      ⚠ Repaired truncated JSON (closed {open_braces} braces)")
            return _normalize_extracted_json(parsed)
    except json.JSONDecodeError:
        pass
    
    # Attempt 3: Remove trailing commas
    try:
        fixed = re.sub(r',\s*([}\]])', r'\1', cleaned)
        return json.loads(fixed)
    except json.JSONDecodeError:
        pass
    
    # Attempt 4: Ask LLM to reformat
    return _try_reformat_response(raw)


# ============================================================
# Path B: PDF Annual Reports (non-listed institutions)
# ============================================================
def collect_pdf_reports(company: dict, query: str, target_years: list, data_dir: str) -> dict:
    """PDF pipeline: Tavily search PDF → download → LLM extract, with missing year retry."""
    search_name = company.get("search_name", company["name"])
    
    # Search for PDF annual reports (filtered by target years)
    pdf_links = _search_pdf_reports(search_name, target_years)
    if not pdf_links:
        print(f"    ✗ No PDF reports found for {search_name}")
        return None
    
    # Download PDFs
    raw_dir = os.path.join(data_dir, "raw", re.sub(r'[^\w]', '_', search_name))
    os.makedirs(raw_dir, exist_ok=True)
    
    pdf_paths = []
    for link in pdf_links:
        path = _download_pdf(link["url"], raw_dir, link.get("filename", "report.pdf"))
        if path:
            pdf_paths.append(path)
            print(f"    ✓ Downloaded: {os.path.basename(path)}")
    
    if not pdf_paths:
        return None
    
    # Extract data from PDFs
    result = _extract_from_pdfs(pdf_paths, query, target_years)
    
    # Check for missing years and retry search for each missing year
    if result and result.get("data"):
        found_years = set()
        for group in result["data"].values():
            for metric in group.values():
                if isinstance(metric, dict):
                    found_years.update(str(y) for y in metric.keys())
        
        missing = [str(y) for y in target_years if str(y) not in found_years]
        if missing:
            print(f"    Missing years: {missing}, searching specifically...")
            for year in missing:
                retry_links = _search_pdf_reports(search_name, [int(year)])
                for link in retry_links:
                    if link.get("url") not in [l.get("url") for l in pdf_links]:
                        path = _download_pdf(link["url"], raw_dir, link.get("filename", "report.pdf"))
                        if path:
                            print(f"    ✓ Retry downloaded: {os.path.basename(path)}")
                            extra = _extract_from_pdfs([path], query, target_years)
                            if extra and extra.get("data"):
                                for g, metrics in extra["data"].items():
                                    if g not in result["data"]:
                                        result["data"][g] = {}
                                    for m, years in metrics.items():
                                        if m not in result["data"][g]:
                                            result["data"][g][m] = {}
                                        for y, v in years.items():
                                            if y not in result["data"][g][m]:
                                                result["data"][g][m][y] = v
    
    return result


# Known institution → official annual report page mapping
# For each entry: url = report listing page, link_selector = CSS selector for PDF links
KNOWN_REPORT_SOURCES = {
    "微众银行": {
        "url": "https://www.webank.com/financialReport/list",
        "base_url": "https://www.webank.com",
        "link_selector": "div.report-list a, div.list-wrap a, a[href*='.pdf']",
    },
    # Add more institutions here as needed:
    # "网商银行": {"url": "https://...", "base_url": "https://...", "link_selector": "..."},
}


def _fetch_official_reports(company_name: str, target_years: list) -> list:
    """Try to get PDF links from official website for known institutions."""
    config = None
    for name, cfg in KNOWN_REPORT_SOURCES.items():
        if name in company_name or company_name in name:
            config = cfg
            break
    
    if not config:
        return []
    
    print(f"    [Official] Fetching from {config['url']}...")
    
    try:
        from bs4 import BeautifulSoup
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'text/html,application/xhtml+xml',
        }
        resp = requests.get(config["url"], headers=headers, timeout=15, verify=False)
        resp.raise_for_status()
        
        soup = BeautifulSoup(resp.text, 'html.parser')
        
        # Try multiple strategies to find PDF links
        results = []
        seen_years = set()
        min_year = str(min(target_years))
        max_year = str(max(target_years))
        
        # Strategy 1: CSS selector from config
        links = soup.select(config["link_selector"])
        
        # Strategy 2: If no results, try all <a> tags with href containing 'pdf' or 'report'
        if not links:
            links = soup.find_all('a', href=True)
        
        for a_tag in links:
            href = a_tag.get('href', '')
            text = a_tag.get_text(strip=True)
            
            # Skip non-PDF and non-report links
            if not href:
                continue
            
            # Make absolute URL
            if href.startswith('/'):
                href = config["base_url"] + href
            elif not href.startswith('http'):
                continue
            
            # Check if it's a report link (PDF or report page)
            is_report = any(kw in (href + text).lower() for kw in ['.pdf', '年报', '年度报告', 'annual', 'report'])
            if not is_report:
                continue
            
            # Extract year
            year_match = re.search(r'20[12]\d', text + href)
            year = year_match.group() if year_match else ""
            
            if year and year not in seen_years and min_year <= year <= max_year:
                seen_years.add(year)
                results.append({
                    "url": href,
                    "title": text or f"{company_name} {year}年度报告",
                    "year": year,
                    "filename": f"{company_name}_{year}_annual_report.pdf",
                    "source": "official_website",
                })
        
        if results:
            print(f"    ✓ Official source: found {len(results)} reports for years {sorted(seen_years, reverse=True)}")
        else:
            print(f"    ⚠ Official page loaded but no report links matched target years")
        
        return results
        
    except Exception as e:
        print(f"    ⚠ Official source failed: {e}")
        return []


def _search_pdf_reports(company_name: str, target_years: list) -> list:
    """Search for PDF annual reports. Priority: official website > Tavily search."""
    
    # Priority 1: Official website (reliable, authoritative)
    official = _fetch_official_reports(company_name, target_years)
    if official:
        return official
    
    # Priority 2: Tavily search (fallback, less reliable)
    print(f"    [Tavily] Searching for PDF reports...")
    min_year = str(min(target_years))
    max_year = str(max(target_years))
    try:
        resp = requests.post("https://api.tavily.com/search", json={
            "api_key": TAVILY_API_KEY,
            "query": f'{company_name} 年度报告 PDF 下载 年报 {min_year}-{max_year}',
            "max_results": len(target_years) + 5,
        }, timeout=30)
        data = resp.json()
        
        results = []
        seen_years = set()
        for r in data.get("results", []):
            url = r.get("url", "")
            title = r.get("title", "")
            # Filter for PDF links
            if url.endswith('.pdf') or 'pdf' in url.lower():
                # Extract year from title or URL
                year_match = re.search(r'20[12]\d', title + url)
                year = year_match.group() if year_match else ""
                # Only include PDFs within target year range
                if year and year not in seen_years and min_year <= year <= max_year:
                    seen_years.add(year)
                    results.append({
                        "url": url,
                        "title": title,
                        "year": year,
                        "filename": f"{company_name}_{year}_annual_report.pdf",
                    })
        print(f"    Found {len(results)} PDF reports in range [{min_year}-{max_year}]: {sorted(seen_years, reverse=True)}")
        return results
    except Exception as e:
        print(f"    ✗ PDF search failed: {e}")
        return []


def _download_pdf(url: str, output_dir: str, filename: str) -> str:
    """Download PDF file. Uses verify=False to bypass SSL cert issues on Windows."""
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    filepath = os.path.join(output_dir, filename)
    if os.path.exists(filepath) and os.path.getsize(filepath) > 10000:
        return filepath
    try:
        resp = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=60, stream=True, verify=False)
        resp.raise_for_status()
        with open(filepath, 'wb') as f:
            for chunk in resp.iter_content(8192):
                f.write(chunk)
        return filepath
    except Exception as e:
        print(f"    ✗ Download failed: {e}")
        return None


def _extract_from_pdfs(pdf_paths: list, query: str, target_years: list) -> dict:
    """Extract data from PDF annual reports with unit normalization."""
    try:
        import pdfplumber
    except ImportError:
        print("    ✗ pdfplumber not installed (pip install pdfplumber)")
        return None
    
    UNIT_TO_QIAN = {
        '千元': 1, '元': 0.001, '万元': 10,
        '百万元': 1000, '亿元': 100000,
        'thousands': 1, 'millions': 1000, 'billions': 1000000,
    }
    
    all_data = {}
    target_unit = "千元"
    
    for pdf_path in pdf_paths:
        print(f"    [PDF Extract] {os.path.basename(pdf_path)}...")
        try:
            text = ""
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages[:80]:
                    tables = page.extract_tables()
                    for table in tables:
                        rows = ['| ' + ' | '.join(str(c or '') for c in row) + ' |' for row in table]
                        text += '\n'.join(rows) + '\n\n'
                    pt = page.extract_text()
                    if pt:
                        text += pt + '\n'
            
            if len(text) > 200000:
                text = text[:200000]
            
            prompt = f"""从以下PDF年报文本中提取结构化数据。

用户需求: {query}
目标年份: {target_years}

必须输出严格的JSON格式如下（不要自定义格式）:
{{
  "data": {{
    "数据组名": {{
      "指标名(English)": {{"2020": 数值, "2021": 数值}}
    }}
  }},
  "metadata": {{"unit": "原始单位如万元/千元/亿元", "notes": ""}}
}}

规则:
- 年份做key，数值做value，不要用数组
- 在metadata.unit中写年报中实际使用的单位
- 不要编造数据
- 只输出JSON，不要输出其他文字

{text[:150000]}"""

            # Retry up to 2 times for PDF extraction
            parsed = None
            for attempt in range(2):
                result = generate_content(prompt=prompt, max_output_tokens=8000)
                if not result:
                    print(f"      ⚠ Attempt {attempt+1}: empty response, retrying...")
                    continue
                try:
                    cleaned = result.strip().strip('```json').strip('```').strip()
                    parsed = json.loads(cleaned)
                    parsed = _normalize_extracted_json(parsed)
                    break
                except Exception as e:
                    print(f"      ⚠ Attempt {attempt+1}: JSON parse failed ({e}), retrying...")
            
            if parsed and parsed.get("data"):
                # Unit normalization
                this_unit = parsed.get("metadata", {}).get("unit", "千元")
                multiplier = 1
                for unit_key, mult in UNIT_TO_QIAN.items():
                    if unit_key in this_unit:
                        multiplier = mult
                        break
                
                for group, metrics in parsed["data"].items():
                    if group not in all_data:
                        all_data[group] = {}
                    for metric, years in metrics.items():
                        if metric not in all_data[group]:
                            all_data[group][metric] = {}
                        for year, val in years.items():
                            if year not in all_data[group][metric]:
                                try:
                                    all_data[group][metric][year] = round(float(val) * multiplier, 2)
                                except (ValueError, TypeError):
                                    all_data[group][metric][year] = val
                
                print(f"      ✓ Extracted {sum(len(m) for m in parsed['data'].values())} metrics")
            else:
                print(f"      ✗ No structured data extracted from this PDF")
        except Exception as e:
            print(f"      ✗ PDF extraction failed: {e}")
    
    return {"data": all_data, "metadata": {"unit": target_unit}} if all_data else None



# ============================================================
# Path B2: Semi-Annual Data (non-listed institutions with 半年报)
# ============================================================
def collect_semi_annual(company: dict, query: str, data_dir: str) -> dict:
    """
    Collect semi-annual (H1/H2) report data for non-listed financial institutions.

    Strategy:
    - Search for 半年报/中期报告 PDFs via Tavily
    - Download and extract with LLM (same pipeline as annual PDF)
    - Period keys use "YYYY-H1" / "YYYY-H2" format

    Used for: Chinese banks, consumer finance companies, insurance companies
    that publish annual reports + semi-annual reports (but not quarterly).
    Generalizable: any company routed to pdf_search path will use this.
    """
    search_name = company.get("search_name", company["name"])
    current_year = datetime.now().year
    target_years = list(range(current_year - 2, current_year + 1))

    print(f"\n  [Semi-annual] Searching for 半年报 PDFs: {search_name}...")
    semi_links = _search_semi_annual_reports(search_name, target_years)

    if not semi_links:
        print(f"    No PDFs found, trying web search for semi-annual data...")
        return _collect_semi_annual_from_web(search_name, query, target_years)

    raw_dir = os.path.join(data_dir, "raw", re.sub(r'[^\w]', '_', search_name) + "_semiannual")
    os.makedirs(raw_dir, exist_ok=True)

    period_paths = []
    for link in semi_links:
        path = _download_pdf(link["url"], raw_dir, link.get("filename", "semi_report.pdf"))
        if path:
            period_paths.append({"path": path, "period": link.get("period", "unknown")})
            print(f"    ✓ Downloaded: {os.path.basename(path)} ({link.get('period', '')})")

    if not period_paths:
        return _collect_semi_annual_from_web(search_name, query, target_years)

    semi_data = {}
    for item in period_paths:
        result = _extract_semi_annual_from_pdf(item["path"], query, item["period"])
        if result and result.get("data"):
            for group, metrics in result["data"].items():
                if group not in semi_data:
                    semi_data[group] = {}
                for metric, periods_data in metrics.items():
                    if metric not in semi_data[group]:
                        semi_data[group][metric] = {}
                    semi_data[group][metric].update(periods_data)
            print(f"    ✓ Extracted: {item['period']}")

    if not semi_data:
        return _collect_semi_annual_from_web(search_name, query, target_years)

    return {"data": semi_data, "metadata": {"source": "pdf_semi_annual", "period_format": "YYYY-H1"}}


def _search_semi_annual_reports(search_name: str, target_years: list) -> list:
    """Search for semi-annual report PDFs via Tavily."""
    min_year = str(min(target_years))
    max_year = str(max(target_years))
    results = []
    seen_periods = set()

    if not TAVILY_API_KEY:
        return []

    try:
        queries = [
            f'{search_name} 半年报 PDF {max_year} {int(max_year)-1}',
            f'{search_name} 中期报告 interim semi-annual report',
        ]
        for q in queries:
            resp = requests.post("https://api.tavily.com/search", json={
                "api_key": TAVILY_API_KEY,
                "query": q,
                "max_results": 5,
            }, timeout=30)
            data = resp.json()
            for r in data.get("results", []):
                url = r.get("url", "")
                title = r.get("title", "")
                if not (url.endswith('.pdf') or 'pdf' in url.lower()):
                    continue
                is_semi = any(kw in (url + title).lower() for kw in ['半年', '中期', 'interim', 'semi'])
                if not is_semi:
                    continue
                year_match = re.search(r'20[12]\d', title + url)
                year = year_match.group() if year_match else ""
                if not year or not (min_year <= year <= max_year):
                    continue
                period = f"{year}-H1"
                if period not in seen_periods:
                    seen_periods.add(period)
                    results.append({
                        "url": url,
                        "title": title,
                        "period": period,
                        "year": year,
                        "filename": f"{search_name}_{year}_半年报.pdf",
                    })
    except Exception as e:
        print(f"    Semi-annual search error: {e}")

    print(f"    Found {len(results)} semi-annual reports: {[r['period'] for r in results]}")
    return results


def _collect_semi_annual_from_web(search_name: str, query: str, target_years: list) -> dict:
    """Fallback: collect semi-annual key metrics via web search when no PDF found."""
    current_year = max(target_years)
    queries = [
        f"{search_name} {current_year}年上半年 净利润 营收 贷款余额",
        f"{search_name} {current_year-1}年上半年 净利润 营收",
        f"{search_name} 半年业绩 H1 {current_year} financial results",
    ]
    all_content = []
    for q in queries:
        if not TAVILY_API_KEY:
            break
        try:
            resp = requests.post("https://api.tavily.com/search", json={
                "api_key": TAVILY_API_KEY,
                "query": q,
                "max_results": 3,
                "include_raw_content": True,
            }, timeout=20)
            data = resp.json()
            for r in data.get("results", []):
                content = r.get("raw_content", r.get("content", ""))[:3000]
                if content:
                    all_content.append(f"Source: {r.get('url', '')}\n{content}")
        except Exception:
            continue

    if not all_content:
        return None

    combined = '\n\n---\n\n'.join(all_content)
    prompt = f"""从以下搜索结果中提取{search_name}的半年度财务数据:

用户需求: {query}

规则:
- 期间用"YYYY-H1"或"YYYY-H2"格式（如"2025-H1"表示2025上半年）
- 只提取明确标注为"上半年"、"H1"、"中期"的数据
- 不编造数据

输出JSON:
{{
  "data": {{
    "数据组名": {{
      "指标名(English)": {{"2025-H1": 值, "2024-H1": 值}}
    }}
  }},
  "metadata": {{"unit": "亿元", "notes": "web search fallback"}}
}}

内容:
{combined}"""

    result = generate_content(prompt=prompt, max_output_tokens=4000)
    if not result:
        return None

    try:
        parsed = json.loads(result.strip().strip('```json').strip('```').strip())
        return _normalize_extracted_json(parsed)
    except Exception:
        return None


def _extract_semi_annual_from_pdf(pdf_path: str, query: str, period_label: str) -> dict:
    """Extract semi-annual data from a PDF 半年报 using same pipeline as annual."""
    try:
        import pdfplumber
    except ImportError:
        return None

    try:
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages[:60]:
                tables = page.extract_tables()
                for table in tables:
                    rows = ['| ' + ' | '.join(str(c or '') for c in row) + ' |' for row in table]
                    text += '\n'.join(rows) + '\n\n'
                pt = page.extract_text()
                if pt:
                    text += pt + '\n'

        if len(text) > 100000:
            text = text[:100000]

        year_match = re.search(r'20[12]\d', period_label)
        year_str = year_match.group() if year_match else str(datetime.now().year - 1)
        prev_year = str(int(year_str) - 1)

        prompt = f"""从以下半年报（中期报告）文本中提取财务数据。

用户需求: {query}
报告期间: {period_label} (上半年/H1)

规则:
- 所有期间key使用"YYYY-H1"格式
  当期用"{year_str}-H1"，去年同期用"{prev_year}-H1"
- 找不到的不编造
- 数值保持原始数字，不转换单位

输出格式:
{{
  "data": {{
    "数据组名": {{
      "指标名(English)": {{"{year_str}-H1": 数值, "{prev_year}-H1": 数值}}
    }}
  }},
  "metadata": {{"unit": "原始单位（如亿元/百万元）", "period": "{period_label}"}}
}}

文本:
{text}"""

        result = generate_content(prompt=prompt, max_output_tokens=6000)
        if not result:
            return None

        parsed = json.loads(result.strip().strip('```json').strip('```').strip())
        return _normalize_extracted_json(parsed)
    except Exception as e:
        print(f"      ✗ Semi-annual PDF extraction failed: {e}")
        return None


def _compute_annual_forecast(company_results: dict) -> dict:
    """
    Compute full-year estimates for companies with partial-year data.

    Logic by data type:
    - Quarterly (SEC 6-K): sum YYYY-Q1~Q4 → actual if 4Q present; extrapolate if <4Q
    - Semi-annual (PDF 半年报): FY estimate = H1 / historical_H1_FY_ratio
    - Neither: skip (e.g. web_search companies without high-frequency data)

    Returns: {company_name: {key: {"value": X, "type": "...", "metric": ..., "group": ...}}}
    """
    forecasts = {}
    current_year = datetime.now().year
    forecast_year = str(current_year - 1)  # e.g. "2025" when running in 2026

    for company_name, result in company_results.items():
        if not result or not result.get("data"):
            continue

        company_fc = {}

        # --- Strategy A: Sum quarterly data (SEC path) ---
        quarterly_groups = {k: v for k, v in result["data"].items() if k.startswith("[季度]")}
        if quarterly_groups:
            for group, metrics in quarterly_groups.items():
                group_key = group.replace("[季度] ", "").replace("[季度]", "")
                for metric, periods in metrics.items():
                    fy_vals = {
                        k: v for k, v in periods.items()
                        if re.match(rf'{forecast_year}-Q\d', str(k)) and isinstance(v, (int, float))
                    }
                    if not fy_vals:
                        continue
                    n = len(fy_vals)
                    total = sum(fy_vals.values())
                    if n == 4:
                        estimate, ftype = total, "实际值 (4季度加总)"
                    elif n >= 2:
                        estimate = round(total * 4 / n)
                        ftype = f"预测值 ({n}季度外推 ×{4/n:.2f})"
                    else:
                        continue
                    fc_key = f"{group_key}|{metric}"
                    company_fc[fc_key] = {
                        "value": estimate, "type": ftype,
                        "quarters": n, "metric": metric, "group": group_key,
                    }

        # --- Strategy B: Double semi-annual data (PDF path) ---
        semi_groups = {k: v for k, v in result["data"].items() if k.startswith("[半年]")}
        if semi_groups and not company_fc:
            annual_groups = {k: v for k, v in result["data"].items()
                             if not k.startswith("[季度]") and not k.startswith("[半年]")}

            for group, metrics in semi_groups.items():
                group_key = group.replace("[半年] ", "").replace("[半年]", "")
                for metric, periods in metrics.items():
                    # Find most recent H1 value
                    h1_entries = sorted(
                        [(k, v) for k, v in periods.items()
                         if 'H1' in str(k) and isinstance(v, (int, float))],
                        reverse=True
                    )
                    if not h1_entries:
                        continue
                    h1_period, h1_val = h1_entries[0]
                    h1_year_m = re.search(r'20[12]\d', str(h1_period))
                    if not h1_year_m:
                        continue
                    h1_year_str = h1_year_m.group()

                    # Try to find historical H1/FY ratio
                    ratio = 0.48  # default: H1 ≈ 48% of FY for Chinese financial institutions
                    for ag, am in annual_groups.items():
                        for am_metric, am_years in am.items():
                            m_short = metric[:6].lower()
                            if m_short in am_metric.lower() or am_metric[:6].lower() in m_short:
                                fy_val = am_years.get(
                                    h1_year_str,
                                    am_years.get(int(h1_year_str) if h1_year_str.isdigit() else h1_year_str)
                                )
                                if fy_val and isinstance(fy_val, (int, float)) and fy_val > 0:
                                    ratio = min(max(h1_val / fy_val, 0.3), 0.7)
                                    break

                    estimate = round(h1_val / ratio)
                    ftype = f"预测值 (H1÷{ratio:.0%}历史季节性)"
                    fc_key = f"{group_key}|{metric}"
                    company_fc[fc_key] = {
                        "value": estimate, "type": ftype, "h1_ratio": ratio,
                        "h1_period": h1_period, "metric": metric, "group": group_key,
                    }

        if company_fc:
            forecasts[company_name] = company_fc

    return forecasts


# ============================================================
# Path C: Web Search Fallback
# ============================================================
def collect_web_search(company: dict, query: str, target_years: list) -> dict:
    """Web search pipeline: multi-query → LLM summarize."""
    search_name = company.get("search_name", company["name"])
    current_year = datetime.now().year
    
    queries = [
        f"{search_name} 年度 收入 净利润 {current_year} 财务数据",
        f"{search_name} 贷款余额 信贷 规模",
        f"{search_name} 财报 营收 利润 历年 变化",
        f"{search_name} financial results revenue profit",
    ]
    
    all_content = []
    for q in queries:
        try:
            resp = requests.post("https://api.tavily.com/search", json={
                "api_key": TAVILY_API_KEY,
                "query": q,
                "max_results": 5,
                "include_raw_content": True,
            }, timeout=30)
            data = resp.json()
            for r in data.get("results", []):
                content = r.get("raw_content", r.get("content", ""))[:5000]
                if content:
                    all_content.append(f"Source: {r.get('url', '')}\n{content}")
        except Exception:
            continue
    
    if not all_content:
        print(f"    ✗ No web results for {search_name}")
        return None
    
    print(f"    Found {len(all_content)} web articles")
    
    combined = '\n\n---\n\n'.join(all_content[:20])
    if len(combined) > 150000:
        combined = combined[:150000]
    
    prompt = f"""从搜索结果中提取{search_name}的财务数据:
用户需求: {query}
目标年份: {target_years}

规则:
- 不同来源冲突时选最权威/最新的
- 找不到的不编造
- 在metadata的notes中标注数据可靠性(high/medium/low)
- 输出严格JSON格式，指标名用"中文(English)"

必须使用以下格式（不要自定义key名）:
{{
  "data": {{
    "数据组名": {{
      "指标名": {{"2020": 值, "2021": 值}}
    }}
  }},
  "metadata": {{"unit": "...", "notes": "...", "reliability": "medium"}}
}}

重要: 即使数据稀少也要用上述格式。年份做key，数值做value。

{combined}"""

    result = generate_content(prompt=prompt, max_output_tokens=8000)
    if not result:
        print(f"    ✗ LLM returned empty response")
        return None
    
    try:
        parsed = json.loads(result.strip().strip('```json').strip('```').strip())
        # Normalize: if LLM used a custom top-level key instead of "data"
        parsed = _normalize_extracted_json(parsed)
        if parsed and parsed.get("data"):
            metrics_count = sum(
                len(yrs) for g in parsed["data"].values()
                for m, yrs in g.items() if isinstance(yrs, dict)
            )
            print(f"    ✓ Extracted {metrics_count} data points (web search)")
        return parsed
    except Exception as e:
        print(f"    ✗ JSON parse failed: {e}")
        # Try to salvage: ask LLM to reformat
        return _try_reformat_response(result)


# ============================================================
# Path D: CN-Listed (reuse collect_financials.py)
# ============================================================
def collect_cn_listed(company: dict, query: str, target_years: list, data_dir: str) -> dict:
    """CN-listed: use existing collect_financials.py."""
    from collect_financials import run_collection
    ticker = company["ticker"]
    output_path = os.path.join(data_dir, f"cn_{ticker}.json")
    result = run_collection(ticker=ticker, output_path=output_path)
    
    if result:
        # Convert to Type 8 format
        data = {}
        if result.get("financials", {}).get("income_statement"):
            data["利润表(Income Statement)"] = {}
            for year, vals in result["financials"]["income_statement"].items():
                for k, v in vals.items():
                    if k != "source":
                        label = k
                        if label not in data["利润表(Income Statement)"]:
                            data["利润表(Income Statement)"][label] = {}
                        data["利润表(Income Statement)"][label][year] = v
        
        if result.get("financials", {}).get("indicators"):
            data["财务指标(Financial Indicators)"] = {}
            for year, vals in result["financials"]["indicators"].items():
                for k, v in vals.items():
                    if k != "source":
                        if k not in data["财务指标(Financial Indicators)"]:
                            data["财务指标(Financial Indicators)"][k] = {}
                        data["财务指标(Financial Indicators)"][k][year] = v
        
        if result.get("market_data"):
            data["市场数据(Market Data)"] = {
                "当前价格(Current Price)": {"latest": result["market_data"].get("price", "")},
            }
        
        return {"data": data, "metadata": {"source": "cn_listed:eastmoney_f10"}} if data else None
    return None


# ============================================================
# Report Generation (Word)
# ============================================================
def _format_number(val) -> str:
    """Format number with thousand separators. Handles int, float, str."""
    if val == "" or val is None:
        return "-"
    try:
        num = float(val)
        if num == int(num):
            return f"{int(num):,}"
        return f"{num:,.2f}"
    except (ValueError, TypeError):
        return str(val)


def _flatten_value(val) -> str:
    """Flatten nested dicts/lists into a readable string for table cells."""
    if isinstance(val, dict):
        # e.g. {'2023': 123, '2024': 456} -> should not appear as cell value
        # But if it does, format as 'k1: v1, k2: v2'
        parts = [f"{k}: {_format_number(v)}" for k, v in val.items()]
        return "; ".join(parts)
    elif isinstance(val, list):
        return ", ".join(str(v) for v in val)
    else:
        return _format_number(val)


def _set_cell_font(cell, size=None, cn_font='SimSun', en_font='Arial'):
    """Set cell font to 宋体+Arial."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    if size is None:
        size = Pt(9)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = size
            run.font.name = en_font
            run.element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)


def generate_word_report(company_results: dict, query: str, years: int, output_dir: str) -> str:
    """Generate Word report with pure data tables. Zero LLM calls."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    
    doc = Document()
    
    # Style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # Title
    title = doc.add_heading('财务数据提取报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Arial'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    doc.add_paragraph(f'查询: {query}')
    doc.add_paragraph(f'时间范围: 近{years}年')
    doc.add_paragraph(f'生成日期: {datetime.now().strftime("%Y-%m-%d")}')
    doc.add_paragraph(f'公司: {", ".join(company_results.keys())}')
    doc.add_paragraph()
    
    # ================================================================
    # Part 1: Annual Data Tables (per company)
    # ================================================================
    doc.add_heading('一、年度数据', level=1)
    
    appendix_letter = ord('A')
    has_quarterly = False
    has_semi_annual = False
    
    for company_name, result in company_results.items():
        if not result or not result.get("data"):
            doc.add_heading(f'{chr(appendix_letter)}. {company_name}', level=2)
            doc.add_paragraph('未获取到有效数据。')
            appendix_letter += 1
            continue
        
        # Split annual vs high-frequency (quarterly / semi-annual)
        annual_groups = {k: v for k, v in result["data"].items()
                         if not k.startswith("[季度]") and not k.startswith("[半年]")}
        quarterly_groups = {k: v for k, v in result["data"].items() if k.startswith("[季度]")}
        semi_annual_groups = {k: v for k, v in result["data"].items() if k.startswith("[半年]")}
        if quarterly_groups:
            has_quarterly = True
        if semi_annual_groups:
            has_semi_annual = True
        
        if not annual_groups:
            appendix_letter += 1
            continue
        
        # Company heading with source tag
        source = result.get("metadata", {}).get("source", "unknown")
        unit = result.get("metadata", {}).get("unit", "")
        source_short = source.split(':')[0].replace('sec_edgar', 'SEC').replace('cn_listed', '东财').replace('pdf_search', 'PDF年报').replace('web_search', 'Web')
        
        heading_text = f'{chr(appendix_letter)}. {company_name}（{source_short}）'
        doc.add_heading(heading_text, level=2)
        appendix_letter += 1
        
        if unit:
            p = doc.add_paragraph(f'单位: {unit}')
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 100, 100)
        
        for group_name, metrics in annual_groups.items():
            doc.add_heading(group_name, level=3)
            
            # Collect all 4-digit year keys
            all_years = sorted(set(
                str(y) for m in metrics.values()
                if isinstance(m, dict)
                for y in m.keys()
                if re.match(r'^\d{4}$', str(y))
            ))
            
            if not all_years:
                continue
            
            # Build table
            headers = ['指标'] + all_years
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                _set_cell_font(cell, size=Pt(9))
                cell.paragraphs[0].runs[0].bold = True
            
            for metric_name, years_data in metrics.items():
                if not isinstance(years_data, dict):
                    continue
                year_keys = [k for k in years_data.keys() if re.match(r'^\d{4}$', str(k))]
                if not year_keys:
                    continue
                
                row = table.add_row()
                row.cells[0].text = metric_name
                _set_cell_font(row.cells[0])
                
                for i, year in enumerate(all_years, 1):
                    val = years_data.get(year, years_data.get(int(year) if year.isdigit() else year, ""))
                    row.cells[i].text = _flatten_value(val)
                    _set_cell_font(row.cells[i])
            
            doc.add_paragraph()  # spacing
    
    # ================================================================
    # Part 2: Quarterly Data (if available)
    # ================================================================
    if has_quarterly:
        doc.add_heading('二、最近季度数据', level=1)
        
        for company_name, result in company_results.items():
            if not result or not result.get("data"):
                continue
            
            quarterly_groups = {k: v for k, v in result["data"].items() if k.startswith("[季度]")}
            if not quarterly_groups:
                continue
            
            doc.add_heading(f'{company_name}', level=2)
            
            for group_name, metrics in quarterly_groups.items():
                display_name = group_name.replace("[季度] ", "").replace("[季度]", "")
                doc.add_heading(display_name, level=3)
                
                all_periods = sorted(set(
                    str(p) for m in metrics.values()
                    if isinstance(m, dict)
                    for p in m.keys()
                    if re.match(r'^\d{4}-Q\d$', str(p))
                ))
                
                if not all_periods:
                    continue
                
                headers = ['指标'] + all_periods
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                
                for i, h in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    _set_cell_font(cell, size=Pt(9))
                    cell.paragraphs[0].runs[0].bold = True
                
                for metric_name, periods_data in metrics.items():
                    if not isinstance(periods_data, dict):
                        continue
                    quarter_keys = [k for k in periods_data.keys() if re.match(r'^\d{4}-Q\d$', str(k))]
                    if not quarter_keys:
                        continue
                    
                    row = table.add_row()
                    row.cells[0].text = metric_name
                    _set_cell_font(row.cells[0])
                    
                    for i, period in enumerate(all_periods, 1):
                        val = periods_data.get(period, "")
                        row.cells[i].text = _flatten_value(val)
                        _set_cell_font(row.cells[i])
                
                doc.add_paragraph()
    
    # ================================================================
    # ================================================================
    # Part 3: Semi-Annual Data (non-listed institutions with 半年报)
    # ================================================================
    if has_semi_annual:
        doc.add_heading('三、最近半年度数据', level=1)
        doc.add_paragraph('数据来源: 年报PDF或公告。期间格式: YYYY-H1 (上半年) / YYYY-H2 (下半年)。')

        for company_name, result in company_results.items():
            if not result or not result.get("data"):
                continue

            semi_groups = {k: v for k, v in result["data"].items() if k.startswith("[半年]")}
            if not semi_groups:
                continue

            doc.add_heading(f'{company_name}', level=2)

            for group_name, metrics in semi_groups.items():
                display_name = group_name.replace("[半年] ", "").replace("[半年]", "")
                doc.add_heading(display_name, level=3)

                all_periods = sorted(set(
                    str(p) for m in metrics.values()
                    if isinstance(m, dict)
                    for p in m.keys()
                    if re.match(r'^\d{4}-H[12]$', str(p))
                ))

                if not all_periods:
                    continue

                headers = ['指标'] + all_periods
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'

                for i, h in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    _set_cell_font(cell, size=Pt(9))
                    cell.paragraphs[0].runs[0].bold = True

                for metric_name, periods_data in metrics.items():
                    if not isinstance(periods_data, dict):
                        continue
                    h_keys = [k for k in periods_data.keys() if re.match(r'^\d{4}-H[12]$', str(k))]
                    if not h_keys:
                        continue

                    row = table.add_row()
                    row.cells[0].text = metric_name
                    _set_cell_font(row.cells[0])

                    for i, period in enumerate(all_periods, 1):
                        val = periods_data.get(period, "")
                        row.cells[i].text = _flatten_value(val)
                        _set_cell_font(row.cells[i])

                doc.add_paragraph()

    # ================================================================
    # Part 4: Full-Year Forecast (from quarterly or semi-annual data)
    # ================================================================
    _forecasts = _compute_annual_forecast(company_results)
    if _forecasts:
        _next_part = 2 + (1 if has_quarterly else 0) + (1 if has_semi_annual else 0) + 1
        _part_zh = ['一', '二', '三', '四', '五', '六'][min(_next_part - 1, 5)]
        doc.add_heading(f'{_part_zh}、全年预测值', level=1)
        _forecast_year = str(datetime.now().year - 1)
        doc.add_paragraph(
            f'基于高频数据（季报/半年报）外推的{_forecast_year}年全年估算值。'
            '实际值（4季度齐全时）标注"实际值"，外推值标注"预测值"及推算方法。'
        )

        for company_name, fc_data in _forecasts.items():
            if not fc_data:
                continue
            doc.add_heading(f'{company_name}', level=2)

            groups_seen = {}
            for key, item in fc_data.items():
                g = item.get("group", "预测数据")
                if g not in groups_seen:
                    groups_seen[g] = []
                groups_seen[g].append(item)

            for group_key, items in groups_seen.items():
                doc.add_heading(group_key, level=3)
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'

                for i, h in enumerate(['指标', f'{_forecast_year}预测/实际值', '推算依据', '备注']):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    _set_cell_font(cell, size=Pt(9))
                    cell.paragraphs[0].runs[0].bold = True

                for item in items:
                    row = table.add_row()
                    row.cells[0].text = item.get("metric", "")
                    row.cells[1].text = _format_number(item.get("value", ""))
                    row.cells[2].text = item.get("type", "")
                    row.cells[3].text = ""
                    for cell in row.cells:
                        _set_cell_font(cell)

                doc.add_paragraph()

    # Part 3: Data Source Description
    # ================================================================
    _source_part_num = 1 + (1 if has_quarterly else 0) + (1 if has_semi_annual else 0) + (1 if _forecasts else 0) + 1
    _source_zh = ['一', '二', '三', '四', '五', '六'][min(_source_part_num - 1, 5)]
    doc.add_heading(f'{_source_zh}、数据来源', level=1)
    
    source_descriptions = {
        'sec_edgar': 'SEC EDGAR（20-F/10-K 表格提取）',
        'cn_listed': '东方财富 F10',
        'pdf_search': '年报 PDF（Tavily 搜索 + LLM 提取）',
        'web_search': '公开网络信息（Tavily 搜索 + LLM 提取）',
    }
    
    for company_name, result in company_results.items():
        if not result:
            continue
        metadata = result.get('metadata', {})
        source_type = metadata.get('source', '').split(':')[0] if metadata.get('source') else ''
        source_desc = source_descriptions.get(source_type, '未知来源')
        unit = metadata.get('unit', '')
        reliability = metadata.get('reliability', '')
        web_supp = metadata.get('web_supplemented_years', [])
        
        line_parts = [f'{company_name}: {source_desc}']
        if unit:
            line_parts.append(f'单位: {unit}')
        if reliability:
            line_parts.append(f'可靠性: {reliability}')
        if web_supp:
            line_parts.append(f'Web补充年份: {", ".join(web_supp)}')
        
        p = doc.add_paragraph(' | '.join(line_parts))
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # Disclaimer
    doc.add_paragraph()
    disclaimer = doc.add_paragraph('免责声明: 数据由自动化程序从公开渠道提取，可能存在提取误差，仅供参考。')
    for run in disclaimer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save
    os.makedirs(output_dir, exist_ok=True)
    filename = f"financial_data_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


# ============================================================
# Main Pipeline
# ============================================================
def run_pipeline(companies: list, query: str, years: int = 5, output_dir: str = None):
    """
    Main entry point: multi-company financial data extraction.
    
    Args:
        companies: List of company names (any language)
        query: What data to extract (e.g., "分产品收入和利润")
        years: How many years of data
        output_dir: Where to save results
    """
    if not output_dir:
        output_dir = os.path.join(SCRIPT_DIR, "..", "data", "type8_results")
    
    data_dir = os.path.join(output_dir, "data")
    os.makedirs(data_dir, exist_ok=True)
    
    current_year = datetime.now().year
    target_years = list(range(current_year - years, current_year + 1))
    
    print("=" * 60)
    print(f"Type 8: Financial Data Extraction")
    print(f"  Companies: {companies}")
    print(f"  Query: {query}")
    print(f"  Years: {target_years}")
    print("=" * 60)
    
    company_results = {}
    
    for name in companies:
        print(f"\n{'─' * 40}")
        print(f"Processing: {name}")
        print(f"{'─' * 40}")
        
        # Dynamic data source detection
        company = detect_data_source(name)
        source = company["source"]
        
        try:
            if source == "sec_edgar":
                result = collect_sec_edgar(company, query, target_years, data_dir)
            elif source == "cn_listed":
                result = collect_cn_listed(company, query, target_years, data_dir)
            elif source == "pdf_search":
                result = collect_pdf_reports(company, query, target_years, data_dir)
            elif source == "web_search":
                result = collect_web_search(company, query, target_years)
            else:
                result = None
            
            # === Primary Source Failure Fallback ===
            # If primary source failed entirely (e.g. SEC 503, PDF download error),
            # fall back to web search so the company doesn't get zero data.
            if not result or not result.get("data"):
                if source != "web_search":
                    print(f"\n  [Fallback] Primary source '{source}' returned no data — trying web search...")
                    result = collect_web_search({"name": name, "search_name": name}, query, target_years)
                    if result and result.get("data"):
                        result["metadata"]["source"] = f"web_search (fallback from {source})"
                        print(f"  ✓ Web search fallback succeeded")
                    else:
                        print(f"  ✗ Web search fallback also failed")
            
            # Ensure metadata has source type for report footer
            if result and isinstance(result, dict):
                if 'metadata' not in result:
                    result['metadata'] = {}
                if 'source' not in result.get('metadata', {}):
                    result['metadata']['source'] = source
            
            company_results[name] = result
            
            # === Latest-Year Web Search Fallback ===
            # If primary source missed recent years (annual report not yet published),
            # automatically supplement with web search data.
            if result and result.get("data") and source != "web_search":
                found_years = set()
                for group in result["data"].values():
                    for metric in group.values():
                        if isinstance(metric, dict):
                            found_years.update(str(y) for y in metric.keys())
                
                # Only supplement recent missing years (current year and last year)
                recent_threshold = current_year - 1  # e.g., 2025 if current is 2026
                recent_missing = [
                    str(y) for y in target_years 
                    if str(y) not in found_years and y >= recent_threshold
                ]
                
                if recent_missing:
                    print(f"\n  [Fallback] Missing recent years {recent_missing} — supplementing with web search...")
                    web_company = {"name": name, "search_name": name}
                    web_result = collect_web_search(web_company, query, [int(y) for y in recent_missing])
                    
                    if web_result and web_result.get("data"):
                        # Merge web data into existing result (don't overwrite existing data)
                        for group, metrics in web_result["data"].items():
                            if group not in result["data"]:
                                result["data"][group] = {}
                            for metric, years in metrics.items():
                                if metric not in result["data"][group]:
                                    result["data"][group][metric] = {}
                                for year, val in years.items():
                                    if year not in result["data"][group][metric]:
                                        result["data"][group][metric][year] = val
                        
                        # Tag the supplemented data source
                        result["metadata"]["web_supplemented_years"] = recent_missing
                        print(f"  ✓ Web search supplemented years: {recent_missing}")
                    else:
                        print(f"  ⚠ Web search found no data for {recent_missing} (annual report not yet available)")
            
            # === Quarterly Data Supplement (SEC path only) ===
            # For SEC-sourced companies, also collect recent 5 quarters
            if source == "sec_edgar" and result and result.get("data"):
                print(f"\n  [Quarterly] Collecting recent 5 quarters for {name}...")
                try:
                    quarterly_result = collect_sec_quarterly(company, query, 5, data_dir)
                    if quarterly_result and quarterly_result.get("data"):
                        # Merge quarterly data into existing result
                        for group, metrics in quarterly_result["data"].items():
                            if group not in result["data"]:
                                result["data"][group] = {}
                            for metric, periods in metrics.items():
                                if metric not in result["data"][group]:
                                    result["data"][group][metric] = {}
                                result["data"][group][metric].update(periods)
                        
                        result["metadata"]["has_quarterly"] = True
                        quarterly_groups = len(quarterly_result["data"])
                        print(f"  ✓ Quarterly data added: {quarterly_groups} groups")
                    else:
                        print(f"  ⚠ No quarterly data available")
                except Exception as e:
                    print(f"  ⚠ Quarterly collection failed: {e}")

            # === Semi-Annual Data (PDF/non-listed path) ===
            # For non-listed institutions with annual reports (banks, consumer finance co.),
            # also collect semi-annual (H1) data as high-frequency proxy.
            # Generalizable: triggers for any company on pdf_search route.
            if source == "pdf_search" and result and result.get("data"):
                print(f"\n  [Semi-annual] Collecting H1 data for {name}...")
                try:
                    semi_result = collect_semi_annual(company, query, data_dir)
                    if semi_result and semi_result.get("data"):
                        for group, metrics in semi_result["data"].items():
                            semi_group = f"[半年] {group}" if not group.startswith("[半年]") else group
                            if semi_group not in result["data"]:
                                result["data"][semi_group] = {}
                            for metric, periods_data in metrics.items():
                                if metric not in result["data"][semi_group]:
                                    result["data"][semi_group][metric] = {}
                                result["data"][semi_group][metric].update(periods_data)
                        result["metadata"]["has_semi_annual"] = True
                        print(f"  ✓ Semi-annual data added: {len(semi_result['data'])} groups")
                    else:
                        print(f"  ⚠ No semi-annual data available for {name}")
                except Exception as e:
                    print(f"  ⚠ Semi-annual collection failed: {e}")

            if result and result.get("data"):
                groups = list(result["data"].keys())
                total_metrics = sum(len(m) for g in result["data"].values() for m in g.values() if isinstance(m, dict))
                print(f"  ✓ Success: {len(groups)} groups, ~{total_metrics} data points")
            else:
                print(f"  ✗ No data extracted")
        
        except Exception as e:
            print(f"  ✗ Pipeline error: {e}")
            company_results[name] = None
    
    # Save raw JSON
    json_path = os.path.join(output_dir, "extracted_data.json")
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(company_results, f, ensure_ascii=False, indent=2, default=str)
    print(f"\n  [JSON] {json_path}")
    
    # Generate raw data Word report (tables only)
    try:
        word_path = generate_word_report(company_results, query, years, output_dir)
        print(f"  [Word/Raw] {word_path}")
    except Exception as e:
        print(f"  [Word/Raw] Failed: {e}")
        word_path = None

    # Generate analysis report (LLM narrative + tables) via generate_analysis_report.py
    try:
        import subprocess as _sp
        _script = os.path.join(os.path.dirname(os.path.abspath(__file__)), "generate_analysis_report.py")
        _result = _sp.run(
            [sys.executable, _script,
             "--data", json_path,
             "--output", output_dir,
             "--query", query,
             "--years", str(years),
             "--name", f"analysis_report_{datetime.now().strftime('%Y%m%d')}"],
            capture_output=True, text=True, timeout=600
        )
        if _result.returncode == 0:
            # Extract output path from script stdout
            _out_lines = [l for l in _result.stdout.splitlines() if "Saved:" in l or "DONE:" in l]
            print(f"  [Word/Analysis] Generated" + (f": {_out_lines[-1].split('DONE:')[-1].strip()}" if _out_lines else ""))
        else:
            print(f"  [Word/Analysis] Failed (exit {_result.returncode}): {_result.stderr[-300:]}")
    except Exception as e:
        print(f"  [Word/Analysis] Failed: {e}")

    print(f"\n{'=' * 60}")
    print("PIPELINE COMPLETE")
    for name, result in company_results.items():
        status = "✓" if result and result.get("data") else "✗"
        source = detect_data_source(name)["source"] if result else "failed"
        print(f"  {status} {name} [{source}]")
    print(f"{'=' * 60}")
    
    return company_results


# ============================================================
# CLI Entry Point
# ============================================================
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Type 8: Financial Data Extraction")
    parser.add_argument("--companies", required=True,
                        help="Comma-separated company names (e.g., '蚂蚁集团,微众银行,奇富科技')")
    parser.add_argument("--query", default="分产品收入和利润",
                        help="What data to extract (e.g., '分产品收入和利润')")
    parser.add_argument("--years", type=int, default=5,
                        help="Number of years (default: 5)")
    parser.add_argument("--output", default=None,
                        help="Output directory")
    parser.add_argument("--mode", choices=["annual", "earnings"], default="annual",
                        help="annual=年报数据提取(default), earnings=季度业绩分析")
    
    args = parser.parse_args()
    companies = [c.strip() for c in args.companies.split(",")]
    
    if args.mode == "earnings":
        # Route to quarterly earnings analysis pipeline
        from collect_earnings import run_earnings_pipeline
        run_earnings_pipeline(
            companies=companies,
            query=args.query,
            output_dir=args.output,
        )
    else:
        # Default: annual data extraction
        run_pipeline(
            companies=companies,
            query=args.query,
            years=args.years,
            output_dir=args.output,
        )
