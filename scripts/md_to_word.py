# -*- coding: utf-8 -*-
"""
Markdown to Word Converter - Standardized for Professional Research Skill

Features:
- Unified fonts: SimSun (Chinese) + Arial (English/Numbers)
- Automatic Markdown syntax cleaning
- Styled tables with dark blue headers
- Support for headings, tables, lists, paragraphs

Usage:
    python md_to_word.py --input report.md --output report.docx
    python md_to_word.py --input report.md  # Output: report.docx (same dir)
"""

import argparse
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Error: python-docx not installed. Run: python -m pip install python-docx")
    sys.exit(1)


# ============================================================
# FONT CONFIGURATION - MODIFY HERE TO CHANGE DEFAULTS
# ============================================================
FONT_CONFIG = {
    'chinese': '宋体',          # 中文字体
    'english': 'Arial',         # 英文/数字字体
    'title_size': 18,           # 标题字号
    'heading1_size': 16,        # 一级标题字号
    'heading2_size': 14,        # 二级标题字号
    'heading3_size': 12,        # 三级标题字号
    'body_size': 11,            # 正文字号
    'table_header_bg': '1B3A5C', # 表格表头背景色 (深蓝)
    'table_header_text': RGBColor(255, 255, 255),  # 表格表头文字颜色 (白色)
}


def set_font(run, size=None, bold=False, color=None):
    """Apply unified font settings to a run."""
    size = size or FONT_CONFIG['body_size']
    run.font.name = FONT_CONFIG['english']
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # Set East Asian font for Chinese characters
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CONFIG['chinese'])


def add_hyperlink(paragraph, text, url, color=None):
    """Add a real hyperlink to a Word paragraph.
    
    python-docx doesn't natively support hyperlinks, so we use OOXML directly.
    """
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"></w:hyperlink>')
    
    run_elem = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>')
    
    # Apply font settings
    rPr = run_elem.find(qn('w:rPr'))
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_CONFIG["english"]}" w:eastAsia="{FONT_CONFIG["chinese"]}" w:hAnsi="{FONT_CONFIG["english"]}"/>')
    sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="{FONT_CONFIG["body_size"] * 2}"/>')
    link_color = color or RGBColor(0x05, 0x63, 0xC1)
    color_elem = parse_xml(f'<w:color {nsdecls("w")} w:val="{link_color}"/>')
    u_elem = parse_xml(f'<w:u {nsdecls("w")} w:val="single"/>')
    rPr.append(rFonts)
    rPr.append(sz)
    rPr.append(color_elem)
    rPr.append(u_elem)
    
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def add_rich_text(paragraph, text, size=None, default_bold=False, default_color=None):
    """Add text with inline Markdown formatting (bold, links) to a paragraph.
    
    Parses **bold**, [text](url) links, and renders them with proper Word formatting.
    Falls back to clean_markdown_syntax() for other syntax.
    """
    # Pattern to match: **bold**, [text](url), or plain text between them
    pattern = re.compile(r'(\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        
        # Bold: **text**
        bold_match = re.match(r'^\*\*(.+)\*\*$', part)
        if bold_match:
            inner = bold_match.group(1)
            # Check if bold text contains a link
            link_in_bold = re.match(r'^\[([^\]]+)\]\(([^)]+)\)$', inner)
            if link_in_bold:
                add_hyperlink(paragraph, link_in_bold.group(1), link_in_bold.group(2))
            else:
                run = paragraph.add_run(inner)
                set_font(run, size=size, bold=True, color=default_color)
            continue
        
        # Link: [text](url)
        link_match = re.match(r'^\[([^\]]+)\]\(([^)]+)\)$', part)
        if link_match:
            add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
            continue
        
        # Plain text — clean remaining markdown syntax
        cleaned = clean_markdown_syntax(part)
        if cleaned:
            run = paragraph.add_run(cleaned)
            set_font(run, size=size, bold=default_bold, color=default_color)


def clean_markdown_syntax(text):
    """Remove all Markdown syntax symbols from text (for non-link content)."""
    # Bold: **text** or __text__
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    # Italic: *text* or _text_
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    # Code: `text`
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Links: [text](url) -> text (url) — preserve URL as text fallback
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', text)
    # Strikethrough: ~~text~~
    text = re.sub(r'~~([^~]+)~~', r'\1', text)
    # HTML line breaks
    text = text.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
    return text.strip()


def repair_markdown_tables(content):
    """Pre-process Markdown content to fix common LLM-generated table malformations.
    
    Fixes:
    - Orphaned lines between table rows (text without | that should be part of a table)
    - Rows with mismatched column counts (pads or truncates to match header)
    - Unclosed ** markers inside cells
    - Trailing garbage like lone ')' or '*' after table rows
    - Missing leading/trailing | in table rows
    """
    lines = content.split('\n')
    repaired = []
    in_table = False
    header_col_count = 0
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Detect table start
        if stripped.startswith('|') and stripped.endswith('|') and not in_table:
            in_table = True
            header_col_count = stripped.count('|') - 1  # pipes between cells
            repaired.append(line)
            i += 1
            continue
        
        if in_table:
            # Table row (starts with |)
            if stripped.startswith('|'):
                # Fix: ensure line ends with |
                if not stripped.endswith('|'):
                    # Remove trailing garbage like ')' or lone '**'
                    stripped = stripped.rstrip(')* ')
                    if not stripped.endswith('|'):
                        stripped += ' |'
                
                # Fix: balance column count
                pipe_count = stripped.count('|') - 1
                if pipe_count > 0 and header_col_count > 0 and pipe_count != header_col_count:
                    cells = stripped.split('|')[1:-1]  # exclude first/last empty
                    # Pad if too few columns
                    while len(cells) < header_col_count:
                        cells.append(' ')
                    # Truncate if too many columns
                    cells = cells[:header_col_count]
                    stripped = '| ' + ' | '.join(c.strip() for c in cells) + ' |'
                
                repaired.append(stripped)
                i += 1
                continue
            
            # Non-table line while in_table: check if it's orphaned table content
            if stripped and not stripped.startswith('#') and not stripped.startswith('>'):
                # Check if next line is a table row — if so, this is orphaned content, skip it
                if i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
                    # Orphaned line between table rows — discard it
                    i += 1
                    continue
                # Otherwise, table has ended
                in_table = False
                header_col_count = 0
            else:
                # Empty line or heading — table ended
                in_table = False
                header_col_count = 0
            
            repaired.append(line)
            i += 1
            continue
        
        # Not in table — pass through
        repaired.append(line)
        i += 1
    
    result = '\n'.join(repaired)
    
    # Global cleanup: fix unclosed ** pairs (orphaned bold markers)
    # Count ** occurrences; if odd, the last one is orphaned
    double_star_count = result.count('**')
    if double_star_count % 2 != 0:
        # Remove the last orphaned **
        last_pos = result.rfind('**')
        result = result[:last_pos] + result[last_pos + 2:]
    
    return result


def parse_table(lines):
    """Parse Markdown table lines into 2D list."""
    rows = []
    for line in lines:
        if not line.strip().startswith('|'):
            continue
        # Skip separator lines like |---|---| or |:---:|:---:| or partial variants
        stripped = line.strip()
        if re.match(r'^\|[\s\-:|]+\|?\s*$', stripped):
            continue
        # Also skip lines that are purely dashes/spaces (no actual data)
        if re.match(r'^[\s\-|:]+$', stripped):
            continue
        # Split and clean cells
        cells = [clean_markdown_syntax(c.strip()) for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    return rows


def _is_numeric_cell(text):
    """Check if cell text looks like a number, percentage, or currency value."""
    t = text.strip().replace(',', '').replace(' ', '')
    if not t or t in ('-', 'N/A', 'n/a', 'null', '—', '–'):
        return False
    return bool(re.match(r'^[\+\-]?[\$¥€£]?\d', t) or
                re.match(r'^[A-Z]{3}\s', t) or  # RMB 1.2B, USD 500M
                t.endswith('%') or t.endswith('bps'))


def _detect_change_color(text):
    """Return green for positive change, red for negative, None otherwise."""
    t = text.strip()
    if re.match(r'^\+\d', t) or re.search(r'增长|增加|上升|提升', t):
        return RGBColor(0x1A, 0x7A, 0x3A)   # Forest green
    if re.match(r'^-\d', t) or re.match(r'^\(.*\)$', t) or re.search(r'下降|减少|下滑', t):
        return RGBColor(0xC0, 0x39, 0x2B)   # Crimson red
    return None


def _set_cell_borders(cell, color='B0B0B0', width='4'):
    """Set thin borders on all sides of a cell (width in eighths of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{width}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{width}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{width}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{width}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def _set_cell_padding(cell, top=40, bottom=40, left=80, right=80):
    """Set cell margins/padding in twips (1pt = 20 twips)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)


def add_styled_table(doc, data):
    """Add a premium styled table with dark blue header, alternating rows, and smart formatting."""
    if not data or not data[0]:
        return None

    num_cols = len(data[0])
    table = doc.add_table(rows=len(data), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Pre-detect which columns are numeric (check data rows, not header)
    numeric_cols = set()
    for row_data in data[1:]:
        for j, cell_text in enumerate(row_data):
            if j < num_cols and _is_numeric_cell(cell_text):
                numeric_cols.add(j)

    ALT_ROW_COLOR = 'F2F4F7'  # Light blue-gray for alternating rows

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                continue
            cell = table.rows[i].cells[j]
            cell.text = cell_text

            # Cell padding
            _set_cell_padding(cell)
            # Thin borders
            _set_cell_borders(cell)

            # Style text
            for para in cell.paragraphs:
                # Alignment: numeric columns right-aligned, first column left, others center
                if i > 0 and j in numeric_cols:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif i == 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif j == 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                for run in para.runs:
                    if i == 0:  # Header row
                        set_font(run, FONT_CONFIG['body_size'], bold=True,
                                color=FONT_CONFIG['table_header_text'])
                    else:
                        # Detect positive/negative change colors
                        change_color = _detect_change_color(cell_text)
                        set_font(run, FONT_CONFIG['body_size'],
                                color=change_color)

            # Background colors
            if i == 0:
                # Header: dark blue
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{FONT_CONFIG["table_header_bg"]}" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)
            elif i % 2 == 0:
                # Alternating rows: light gray
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{ALT_ROW_COLOR}" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)

    return table


def parse_markdown(content):
    """Parse Markdown content into structured elements."""
    lines = content.split('\n')
    elements = []
    i = 0

    while i < len(lines):
        old_i = i  # Safety: detect if no handler advances i
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Title (# )
        if line.startswith('# ') and not line.startswith('## '):
            elements.append(('title', clean_markdown_syntax(line[2:].strip())))
            i += 1
            continue

        # Heading 1-6
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = clean_markdown_syntax(heading_match.group(2))
            elements.append(('heading', level, text))
            i += 1
            continue

        # Table
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and (lines[i].strip().startswith('|') or
                                      re.match(r'^\|[\s\-:]+\|$', lines[i].strip())):
                table_lines.append(lines[i])
                i += 1
            table_data = parse_table(table_lines)
            if table_data:
                elements.append(('table', table_data))
            continue

        # Blockquote
        if line.strip().startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            elements.append(('blockquote', ' '.join(quote_lines)))
            continue

        # Unordered list
        if re.match(r'^[\-\*\+]\s+', line.strip()):
            list_items = []
            while i < len(lines):
                item_match = re.match(r'^[\-\*\+]\s+(.+)$', lines[i].strip())
                if item_match:
                    list_items.append(item_match.group(1))
                    i += 1
                else:
                    break
            elements.append(('ulist', list_items))
            continue

        # Ordered list
        if re.match(r'^\d+\.\s+', line.strip()):
            list_items = []
            while i < len(lines):
                item_match = re.match(r'^\d+\.\s+(.+)$', lines[i].strip())
                if item_match:
                    list_items.append(item_match.group(1))
                    i += 1
                else:
                    break
            elements.append(('olist', list_items))
            continue

        # Horizontal rule
        if re.match(r'^[\-\*_]{3,}$', line.strip()):
            elements.append(('hr', None))
            i += 1
            continue

        # Paragraph (collect consecutive non-empty lines)
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].startswith(('#', '|', '>', '-', '*', '+')):
            # Check for list start
            if re.match(r'^\d+\.\s+', lines[i].strip()):
                break
            para_lines.append(lines[i])
            i += 1

        if para_lines:
            text = ' '.join(para_lines)
            if text:
                elements.append(('paragraph', text))

        # Safety net: if no handler advanced i, skip this line to prevent infinite loop
        if i == old_i:
            i += 1

    return elements


def build_document(elements):
    """Build Word document from parsed elements."""
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = FONT_CONFIG['english']
    style.font.size = Pt(FONT_CONFIG['body_size'])
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CONFIG['chinese'])

    for elem in elements:
        elem_type = elem[0]

        if elem_type == 'title':
            p = doc.add_heading(elem[1], 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_font(run, FONT_CONFIG['title_size'], bold=True)

        elif elem_type == 'heading':
            level = elem[1]
            text = elem[2]
            # Map markdown heading level to Word heading level (max 3)
            word_level = min(level, 3)
            h = doc.add_heading(text, level=word_level)

            size_map = {1: FONT_CONFIG['heading1_size'],
                       2: FONT_CONFIG['heading2_size'],
                       3: FONT_CONFIG['heading3_size']}
            for run in h.runs:
                set_font(run, size_map.get(word_level, FONT_CONFIG['body_size']), bold=True)

        elif elem_type == 'table':
            add_styled_table(doc, elem[1])
            doc.add_paragraph()  # Space after table

        elif elem_type == 'paragraph':
            p = doc.add_paragraph()
            add_rich_text(p, elem[1])

        elif elem_type == 'ulist':
            for item in elem[1]:
                p = doc.add_paragraph(style='List Bullet')
                add_rich_text(p, item)

        elif elem_type == 'olist':
            for item in elem[1]:
                p = doc.add_paragraph(item, style='List Number')
                for run in p.runs:
                    set_font(run)

        elif elem_type == 'blockquote':
            p = doc.add_paragraph(elem[1])
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                set_font(run)
                run.italic = True

        elif elem_type == 'hr':
            # Add a blank line as separator
            doc.add_paragraph()

    return doc


def convert_md_to_word(input_path, output_path=None):
    """Main conversion function."""
    input_path = Path(input_path)

    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    # Default output path
    if output_path is None:
        output_path = input_path.with_suffix('.docx')
    else:
        output_path = Path(output_path)

    # Read Markdown
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Pre-process: repair malformed tables from LLM output
    content = repair_markdown_tables(content)

    # Parse and build
    elements = parse_markdown(content)
    doc = build_document(elements)

    # Save
    doc.save(str(output_path))

    return output_path


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown to Word with unified formatting (SimSun + Arial)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
    python md_to_word.py --input report.md
    python md_to_word.py --input report.md --output /path/to/output.docx
    python md_to_word.py report.md  # shorthand
        '''
    )
    parser.add_argument('input', nargs='?', help='Input Markdown file path')
    parser.add_argument('--input', '-i', dest='input_file', help='Input Markdown file path')
    parser.add_argument('--output', '-o', help='Output Word file path (default: same dir, .docx extension)')

    args = parser.parse_args()

    # Handle input argument
    input_path = args.input or args.input_file
    if not input_path:
        parser.print_help()
        print("\nError: Input file is required")
        sys.exit(1)

    try:
        output_path = convert_md_to_word(input_path, args.output)
        print(f"✅ Conversion complete!", flush=True)
        print(f"   Input:  {input_path}", flush=True)
        print(f"   Output: {output_path}", flush=True)
        print(f"   Size:   {os.path.getsize(str(output_path)):,} bytes", flush=True)
    except Exception as e:
        print(f"❌ Error: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
